#*TO*#
__pgm__ = "detfull.py"
__version__ = "0.0.0"
__date__ = "01/01/1970 00:00"

"""
Exécution complète de tous les modules de test en mode batch.

Lance chaque module de détection/recherche sur un fichier CSV de test
et génère un rapport DOCX synthétisant les résultats.

NOUVEAUTÉ V2.0 : ÉVALUATION INTELLIGENTE DES COMMENTAIRES
  - Si le CSV de test contient une colonne 'commentaire', chaque commentaire
    est évalué via l'API Anthropic contre les résultats du module detall.
  - 3 statuts possibles : OK, KO, ACTION
  - OK = Le problème décrit semble résolu par la détection actuelle
  - KO = Le problème persiste
  - ACTION = Vérification manuelle nécessaire (instructions fournies)
  - Sans commentaire = comportement inchangé (succès = OK)

MODULES TESTÉS (par ordre alphabétique, configurable via communb.csv section detfull) :

  Détection standard (sans base) :
    detage      : Détection des critères d'âge et de sexe
    detall      : Pipeline complet de détection standard
    detangles   : Détection des angles céphalométriques
    detcount    : Détection LIST/COUNT
    detid       : Détection des identifiants patients
    detmeme     : Détection des recherches par similarité
    dettags     : Détection des pathologies (tags + adjectifs)

  Détection IA (sans base) :
    detia       : Détection IA avec référentiels complets
    detiabrut   : Détection IA brute (sans référentiels, défaut=none)

  Recherche complète (avec base) — Triple détection :
    trouve           : Recherche standard (detall → SQL)
    trouve_ia        : Recherche IA (detia → SQL)
    trouve_iabrut    : Recherche IA brute (detiabrut none → SQL)
    trouveid         : Recherche avec résolution meme+id
    search           : Pipeline search standard (avec fallback IA auto)
    search_purstandard : Pipeline search purstandard (detall seul)
    search_puria     : Pipeline search puria (detia seul)

  Non inclus par défaut :
    detadjs     : Format d'entrée spécial (tag;question)

RAPPORTS DOCX (V3.0) :
  Deux rapports distincts dans tests/ :

  1. {stem}_init.docx  — Rapport initial (AVANT corrections)
     Généré si {stem}_init.docx n'existe PAS encore dans tests/.
     Rapport complet : légende, évaluation commentaires, résultats par module, synthèse.

  2. {stem}_final.docx — Rapport final (APRÈS corrections)
     Généré si {stem}_init.docx EXISTE déjà dans tests/.
     Rapport compact (~5 pages) avec :
       - Partie 1 : Résumé des bugs corrigés (depuis tests/bugs{stem}.docx)
       - Partie 2 : Résumé des résultats après corrections
     Nécessite tests/bugs{stem}.docx (auto-détection).

SYNTAXE CLI :
    python detfull.py <fichier.csv> <base.db> [options]
    python detfull.py <base.db> <fichier.csv> [options]

    Arguments (ordre libre, TOUS DEUX OBLIGATOIRES) :
      *.db     Base SQLite (cherche aussi dans bases/)
      *.csv    Fichier de test (cherche aussi dans tests/)

    Options :
      -v       Mode verbose
      -d       Mode debug
      --no-ia  Exclure les modules IA (detia, detiabrut, trouve_ia, etc.)
      --no-db  Exclure les modules nécessitant une base
      --no-doc Pas de génération du rapport DOCX
      --no-eval Pas d'évaluation IA des commentaires
      --only=x Exécuter uniquement le module x (ex: --only=detage)

EXEMPLES :
    python detfull.py quentin.csv base1964.db
    python detfull.py base1964.db quentin.csv -v
    python detfull.py quentin.csv base1964.db --no-ia
    python detfull.py quentin.csv base1964.db --only=trouve_ia
    python detfull.py quentin.csv base1964.db --no-eval
"""

import os
import sys
import csv
import time
import json
import subprocess
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# ═══════════════════════════════════════════════════════════════════════════
# CONSTANTES
# ═══════════════════════════════════════════════════════════════════════════

# Modules nécessitant une base de données
MODULES_DB = {'trouve', 'trouveid', 'search',
              'trouve_ia', 'trouve_iabrut',
              'search_purstandard', 'search_puria'}

# Modules IA
MODULES_IA = {'detia', 'detiabrut',
              'trouve_ia', 'trouve_iabrut',
              'search_puria'}

# Liste par défaut si communb.csv non disponible
MODULES_DEFAUT = [
    'detage', 'detall', 'detangles', 'detcount', 'detid',
    'detmeme', 'dettags', 'detia', 'detiabrut',
    'trouve', 'trouve_ia', 'trouve_iabrut',
    'trouveid',
    'search', 'search_purstandard', 'search_puria'
]

# Descriptions des modules pour le rapport
DESCRIPTIONS_MODULES = {
    'detage':    "Détecte l'âge et le sexe des patients recherchés",
    'detall':    "Pipeline complet : détecte tout (tags, âge, angles, meme, id...)",
    'detangles': "Détecte les angles céphalométriques (ANB, SNA, SNB)",
    'detcount':  "Détecte si on veut compter (COUNT) ou lister (LIST)",
    'detid':     "Détecte les identifiants patients (id 10122)",
    'detmeme':   "Détecte les recherches par similarité (même portrait que...)",
    'dettags':   "Détecte les pathologies et leurs adjectifs",
    'detia':     "Détection par intelligence artificielle avec référentiels complets",
    'detiabrut': "Détection par intelligence artificielle sans aide (mode brut)",
    'trouve':           "Recherche standard : detall → jsonsql → lancesql",
    'trouve_ia':        "Recherche IA : detia → jsonsql → lancesql",
    'trouve_iabrut':    "Recherche IA brute : detiabrut (none) → jsonsql → lancesql",
    'trouveid':  "Recherche avec résolution des similarités et identifiants",
    'search':           "Pipeline search standard (avec fallback IA automatique)",
    'search_purstandard': "Pipeline search purstandard (detall seul, SANS fallback)",
    'search_puria':     "Pipeline search puria (detia seul, SANS fallback)",
}

# Catégories pour le rapport
CATEGORIES = {
    'standard': {
        'titre': 'Détection standard (sans IA)',
        'emoji': '🔍',
        'description': "Ces modules analysent la question mot par mot avec des règles précises.",
        'modules': ['detcount', 'detid', 'detage', 'detangles', 'dettags', 'detmeme', 'detall'],
    },
    'ia': {
        'titre': 'Détection par Intelligence Artificielle',
        'emoji': '🤖',
        'description': "Ces modules envoient la question à une IA qui comprend le langage naturel.",
        'modules': ['detia', 'detiabrut'],
    },
    'recherche': {
        'titre': 'Recherche dans la base de patients',
        'emoji': '🏥',
        'description': "Ces modules cherchent vraiment les patients dans la base de données.",
        'modules': ['trouve', 'trouve_ia', 'trouve_iabrut', 'trouveid',
                     'search', 'search_purstandard', 'search_puria'],
    },
}

# Mapping modules virtuels → commande réelle + arguments
MODULES_VIRTUELS = {
    'trouve_ia':        ('trouve', ['ia'],       '_ia'),
    'trouve_iabrut':    ('trouve', ['iabrut'],   '_iabrut'),
    'search_purstandard': ('search', ['--mode=purstandard'], '_purstandard'),
    'search_puria':     ('search', ['--mode=puria'],   '_puria'),
}

# Module de référence pour l'évaluation des commentaires
MODULE_EVALUATION = 'detall'


# ═══════════════════════════════════════════════════════════════════════════
# LECTURE CONFIGURATION COMMUNB.CSV
# ═══════════════════════════════════════════════════════════════════════════

def charger_modules_config(verbose=False, debug=False) -> List[str]:
    """
    Charge la liste des modules depuis communb.csv section detfull.
    
    Retourne la liste par défaut si le fichier est introuvable.
    """
    script_dir = Path(__file__).parent
    
    chemins = [
        script_dir / "refs" / "communb.csv",
        script_dir / "refs" / "commun.csv",
        Path("c:/g/refs/communb.csv"),
        Path("c:/cx/refs/communb.csv"),
    ]
    
    for chemin in chemins:
        if chemin.exists():
            if verbose:
                print(f"  Config: {chemin.resolve()}")
            try:
                with open(chemin, 'r', encoding='utf-8-sig') as f:
                    reader = csv.reader(f, delimiter=';')
                    for row in reader:
                        if not row or len(row) < 3:
                            continue
                        if row[0].strip().startswith('#'):
                            continue
                        if row[0].strip() == 'detfull' and row[1].strip() == 'modules':
                            modules = [m.strip() for m in row[2].split(',') if m.strip()]
                            if verbose:
                                print(f"  Modules configurés: {', '.join(modules)}")
                            return modules
            except Exception as e:
                if debug:
                    print(f"[DEBUG] Erreur lecture {chemin}: {e}")
    
    if verbose:
        print(f"  Section detfull non trouvée, utilisation de la liste par défaut")
    return MODULES_DEFAUT.copy()


# ═══════════════════════════════════════════════════════════════════════════
# RÉSOLUTION DE FICHIERS
# ═══════════════════════════════════════════════════════════════════════════

def _trouver_fichier(nom: str, sous_reps: List[str]) -> Optional[Path]:
    """Cherche un fichier dans le répertoire courant puis dans des sous-répertoires."""
    chemin = Path(nom)
    if chemin.exists():
        return chemin.resolve()
    
    script_dir = Path(__file__).parent
    for sous_rep in sous_reps:
        candidat = script_dir / sous_rep / chemin.name
        if candidat.exists():
            return candidat.resolve()
    
    return None


def _trouver_csv(nom: str) -> Optional[Path]:
    """Cherche un fichier CSV dans tests/ et le répertoire courant."""
    return _trouver_fichier(nom, ['tests', '.'])


def _trouver_db(nom: str) -> Optional[Path]:
    """Cherche une base .db dans bases/ et le répertoire courant."""
    return _trouver_fichier(nom, ['bases', '.'])


# ═══════════════════════════════════════════════════════════════════════════
# LECTURE DES QUESTIONS ET COMMENTAIRES DU CSV
# ═══════════════════════════════════════════════════════════════════════════

def lire_questions_csv(fichier_csv: Path) -> List[str]:
    """Lit les questions du fichier CSV (colonne 'question')."""
    questions = []
    try:
        with open(fichier_csv, 'r', encoding='utf-8-sig') as f:
            reader = csv.DictReader(
                (line for line in f if not line.strip().startswith('#')),
                delimiter=';'
            )
            for row in reader:
                q = row.get('question', '').strip()
                if q:
                    questions.append(q)
    except Exception as e:
        print(f"  ⚠ Erreur lecture {fichier_csv}: {e}")
    return questions


def lire_commentaires_csv(fichier_csv: Path) -> List[str]:
    """
    Lit les commentaires du fichier CSV (colonne 'commentaire').
    
    Retourne une liste alignée avec les questions :
    - '' si pas de commentaire pour cette question
    - Le texte du commentaire sinon
    
    Si la colonne 'commentaire' n'existe pas, retourne une liste vide.
    """
    commentaires = []
    try:
        with open(fichier_csv, 'r', encoding='utf-8-sig') as f:
            reader = csv.DictReader(
                (line for line in f if not line.strip().startswith('#')),
                delimiter=';'
            )
            
            # Vérifier si la colonne commentaire existe
            if reader.fieldnames and 'commentaire' not in [f.strip().lower() for f in reader.fieldnames]:
                return []
            
            # Trouver le vrai nom de la colonne (casse insensible)
            col_commentaire = None
            if reader.fieldnames:
                for f in reader.fieldnames:
                    if f.strip().lower() == 'commentaire':
                        col_commentaire = f
                        break
            
            if not col_commentaire:
                return []
            
            for row in reader:
                q = row.get('question', '').strip()
                if q:
                    commentaires.append((row.get(col_commentaire) or '').strip())
    except Exception as e:
        print(f"  ⚠ Erreur lecture commentaires {fichier_csv}: {e}")
    
    return commentaires


# ═══════════════════════════════════════════════════════════════════════════
# LECTURE DES RÉSULTATS CSV (format question;L1;L2;...;Ln)
# ═══════════════════════════════════════════════════════════════════════════

def lire_resultats_csv(fichier_csv: Path) -> List[Dict]:
    """
    Lit un fichier de résultats au format question;L1;L2;...;Ln.
    
    Returns:
        Liste de dicts {'question': str, 'lignes': [str, ...]}
    """
    resultats = []
    if not fichier_csv.exists():
        return resultats
    
    try:
        with open(fichier_csv, 'r', encoding='utf-8-sig') as f:
            reader = csv.reader(
                (line for line in f if not line.strip().startswith('#')),
                delimiter=';'
            )
            entete = next(reader, None)
            if not entete:
                return resultats
            
            idx_question = None
            idx_l_start = None
            for i, col in enumerate(entete):
                col_lower = col.strip().lower()
                if col_lower == 'question':
                    idx_question = i
                elif col_lower == 'l1' and idx_l_start is None:
                    idx_l_start = i
            
            if idx_question is None:
                return resultats
            if idx_l_start is None:
                idx_l_start = 1
            
            for row in reader:
                if len(row) <= idx_question:
                    continue
                question = row[idx_question].strip()
                lignes = [row[i].strip() for i in range(idx_l_start, len(row)) if i < len(row) and row[i].strip()]
                resultats.append({'question': question, 'lignes': lignes})
    except Exception as e:
        print(f"  ⚠ Erreur lecture résultats {fichier_csv}: {e}")
    
    return resultats


# ═══════════════════════════════════════════════════════════════════════════
# ÉVALUATION DES COMMENTAIRES VIA API ANTHROPIC (V2.0)
# ═══════════════════════════════════════════════════════════════════════════

def _installer_anthropic():
    """Installe le package anthropic si nécessaire."""
    try:
        import anthropic
        return True
    except ImportError:
        print("  Installation du package anthropic...")
        proc = subprocess.run(
            [sys.executable, '-m', 'pip', 'install', 'anthropic', '--break-system-packages', '-q'],
            capture_output=True
        )
        if proc.returncode == 0:
            return True
        print(f"  ⚠ Impossible d'installer anthropic: {proc.stderr[:200]}")
        return False


def evaluer_commentaire_ia(question: str, commentaire: str,
                           resultats_detall: List[str],
                           resultats_autres: Dict[str, List[str]] = None,
                           verbose=False, debug=False) -> Dict:
    """
    Évalue un commentaire de testeur contre les résultats des modules via l'API Anthropic.
    
    Args:
        question: Question testée
        commentaire: Commentaire du testeur (observation/attente)
        resultats_detall: Lignes de résultat du module detall pour cette question
        resultats_autres: Dict {module: [lignes]} des autres modules (optionnel)
        verbose: Mode verbose
        debug: Mode debug
    
    Returns:
        dict: {
            'statut': 'OK' | 'KO' | 'ACTION',
            'explication': str,
            'action': str | None  (instructions pour le testeur si ACTION)
        }
    """
    try:
        import anthropic
    except ImportError:
        return {
            'statut': 'ACTION',
            'explication': "Package anthropic non disponible, évaluation impossible",
            'action': "Installer le package anthropic (pip install anthropic)"
        }
    
    # Vérifier la clé API
    if not os.environ.get('ANTHROPIC_API_KEY'):
        return {
            'statut': 'ACTION',
            'explication': "Clé API Anthropic non configurée",
            'action': "Définir la variable d'environnement ANTHROPIC_API_KEY"
        }
    
    # Construire le contexte des résultats
    contexte = "Résultats detall (pipeline complet) :\n"
    if resultats_detall:
        for ligne in resultats_detall:
            contexte += f"  {ligne}\n"
    else:
        contexte += "  (aucun résultat disponible)\n"
    
    if resultats_autres:
        for module, lignes in resultats_autres.items():
            if lignes:
                contexte += f"\nRésultats {module} :\n"
                for ligne in lignes:
                    contexte += f"  {ligne}\n"
    
    prompt = f"""Tu es un évaluateur de tests pour KitView Search, un système de recherche dans une base de patients orthodontiques.

QUESTION TESTÉE :
"{question}"

COMMENTAIRE DU TESTEUR (décrit un problème observé ou une attente) :
"{commentaire}"

RÉSULTATS DE LA DÉTECTION :
{contexte}

CONSIGNE :
Analyse si les résultats de la détection répondent correctement au problème ou à l'attente décrit dans le commentaire du testeur.

Réponds UNIQUEMENT par un JSON valide (sans markdown, sans backticks, sans texte avant/après) :
{{"statut": "OK ou KO ou ACTION", "explication": "Explication courte 1-2 phrases", "action": "Si ACTION: ce que le testeur doit faire manuellement. Si OK ou KO: null"}}

Règles strictes :
- OK = Le résultat de détection montre que le problème décrit dans le commentaire est résolu. Les critères attendus sont bien détectés.
- KO = Le problème persiste. La détection ne gère pas correctement ce cas. Le commentaire décrit un défaut toujours présent.
- ACTION = Impossible de déterminer automatiquement. Le commentaire demande une vérification qui nécessite un accès à la base de données, une comparaison visuelle, ou un jugement humain.

Sois rigoureux : si le commentaire dit qu'un âge n'est pas détecté et que le résultat montre un critère [age], c'est OK.
Si le commentaire signale un problème et que le résultat ne montre aucune détection correspondante, c'est KO."""

    try:
        client = anthropic.Anthropic()
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=500,
            messages=[{"role": "user", "content": prompt}]
        )
        
        texte_reponse = response.content[0].text.strip()
        
        if debug:
            print(f"[DEBUG] Réponse API brute: {texte_reponse[:300]}")
        
        # Nettoyer : enlever d'éventuels backticks markdown
        texte_reponse = texte_reponse.strip()
        if texte_reponse.startswith('```'):
            texte_reponse = texte_reponse.split('\n', 1)[-1]
            if texte_reponse.endswith('```'):
                texte_reponse = texte_reponse[:-3]
            texte_reponse = texte_reponse.strip()
        
        resultat = json.loads(texte_reponse)
        
        # Valider le statut
        statut = resultat.get('statut', 'ACTION').upper()
        if statut not in ('OK', 'KO', 'ACTION'):
            statut = 'ACTION'
        
        return {
            'statut': statut,
            'explication': resultat.get('explication', ''),
            'action': resultat.get('action', None)
        }
    
    except json.JSONDecodeError as e:
        if debug:
            print(f"[DEBUG] Erreur parsing JSON: {e}")
        return {
            'statut': 'ACTION',
            'explication': f"Réponse API non parsable: {texte_reponse[:100]}",
            'action': "Relancer l'évaluation ou vérifier manuellement"
        }
    except Exception as e:
        return {
            'statut': 'ACTION',
            'explication': f"Erreur API: {str(e)[:150]}",
            'action': "Vérifier la connexion et la clé API"
        }


def evaluer_tous_commentaires(questions: List[str], commentaires: List[str],
                              resultats_modules: Dict,
                              verbose=False, debug=False) -> List[Dict]:
    """
    Évalue tous les commentaires des questions contre les résultats des modules.
    
    Args:
        questions: Liste des questions
        commentaires: Liste des commentaires (alignée avec questions)
        resultats_modules: Dict {module: résultat_executer_module}
        verbose: Mode verbose
        debug: Mode debug
    
    Returns:
        Liste de dicts par question :
        [
            {
                'question': str,
                'commentaire': str,
                'statut': 'OK' | 'KO' | 'ACTION' | 'PASS',
                'explication': str,
                'action': str | None
            },
            ...
        ]
    """
    evaluations = []
    
    if not commentaires:
        # Pas de commentaires → toutes les questions sont PASS
        for q in questions:
            evaluations.append({
                'question': q,
                'commentaire': '',
                'statut': 'PASS',
                'explication': 'Pas de commentaire, test passé si module OK',
                'action': None
            })
        return evaluations
    
    # Charger les résultats detall pour la comparaison
    resultats_detall_csv = []
    r_detall = resultats_modules.get('detall')
    if r_detall and r_detall.get('succes') and r_detall.get('fichier_sortie'):
        fichier_detall = Path(r_detall['fichier_sortie'])
        if fichier_detall.exists():
            resultats_detall_csv = lire_resultats_csv(fichier_detall)
    
    # Charger aussi les résultats des modules clés pour contexte
    modules_contexte = ['detage', 'detangles', 'detcount', 'dettags', 'detmeme']
    resultats_contexte_par_question = {}
    for module in modules_contexte:
        r = resultats_modules.get(module)
        if r and r.get('succes') and r.get('fichier_sortie'):
            fichier = Path(r['fichier_sortie'])
            if fichier.exists():
                res_csv = lire_resultats_csv(fichier)
                for idx, res in enumerate(res_csv):
                    if idx not in resultats_contexte_par_question:
                        resultats_contexte_par_question[idx] = {}
                    resultats_contexte_par_question[idx][module] = res.get('lignes', [])
    
    # Évaluer chaque question
    nb_avec_commentaire = sum(1 for c in commentaires if c)
    idx_eval = 0
    
    for i, question in enumerate(questions):
        commentaire = commentaires[i] if i < len(commentaires) else ''
        
        if not commentaire:
            # Pas de commentaire → PASS (l'ancien comportement suffit)
            evaluations.append({
                'question': question,
                'commentaire': '',
                'statut': 'PASS',
                'explication': 'Pas de commentaire, test passé si module OK',
                'action': None
            })
            continue
        
        idx_eval += 1
        print(f"  [{idx_eval}/{nb_avec_commentaire}] Évaluation IA: \"{question[:60]}...\"")
        print(f"        Commentaire: \"{commentaire[:80]}\"")
        
        # Récupérer les résultats detall pour cette question
        lignes_detall = []
        if i < len(resultats_detall_csv):
            lignes_detall = resultats_detall_csv[i].get('lignes', [])
        
        # Récupérer les résultats des autres modules
        autres = resultats_contexte_par_question.get(i, {})
        
        # Appeler l'API
        evaluation = evaluer_commentaire_ia(
            question, commentaire, lignes_detall, autres,
            verbose=verbose, debug=debug
        )
        
        evaluation['question'] = question
        evaluation['commentaire'] = commentaire
        evaluations.append(evaluation)
        
        # Afficher le résultat
        statut = evaluation['statut']
        emoji = {'OK': '✅', 'KO': '❌', 'ACTION': '🔧'}.get(statut, '❓')
        print(f"        → {emoji} {statut} : {evaluation.get('explication', '')[:80]}")
        if evaluation.get('action'):
            print(f"        📋 Action: {evaluation['action'][:80]}")
        print()
    
    return evaluations


# ═══════════════════════════════════════════════════════════════════════════
# EXÉCUTION D'UN MODULE
# ═══════════════════════════════════════════════════════════════════════════

def executer_module(module: str, csv_path: Path, db_path: Optional[Path],
                    verbose: bool = False, debug: bool = False) -> Dict:
    """
    Exécute un module de test en subprocess.
    
    Supporte les modules virtuels (ex: trouve_ia → python trouve.py ... ia)
    via le mapping MODULES_VIRTUELS.
    
    Returns:
        Dict avec: module, fichier_sortie, duree_ms, succes, erreur, nb_questions
    """
    resultat = {
        'module': module,
        'fichier_sortie': None,
        'duree_ms': 0,
        'succes': False,
        'erreur': None,
        'nb_questions': 0,
    }
    
    # Résoudre le module virtuel
    if module in MODULES_VIRTUELS:
        script_reel, extra_args, suffix = MODULES_VIRTUELS[module]
        module_base_db = script_reel
    else:
        script_reel = module
        extra_args = []
        suffix = ''
        module_base_db = module
    
    # Construire la commande
    cmd = [sys.executable, f"{script_reel}.py"]
    
    # Pour search, la base va AVANT le CSV
    if script_reel == 'search':
        if module_base_db in MODULES_DB and db_path:
            cmd.append(str(db_path))
        elif module_base_db in MODULES_DB and not db_path:
            resultat['erreur'] = "Base requise mais non fournie"
            return resultat
        cmd.append(str(csv_path))
    else:
        cmd.append(str(csv_path))
        if module_base_db in MODULES_DB and db_path:
            cmd.append(str(db_path))
        elif module_base_db in MODULES_DB and not db_path:
            resultat['erreur'] = "Base requise mais non fournie"
            return resultat
    
    cmd.extend(extra_args)
    
    if verbose:
        cmd.append('-v')
    if debug:
        cmd.append('-d')
    
    # Fichier de sortie attendu
    stem = csv_path.stem
    fichier_sortie = csv_path.parent / f"{stem}{script_reel}{suffix}.csv"
    resultat['fichier_sortie'] = fichier_sortie
    
    print(f"  {'─' * 60}")
    print(f"  ▶ {module} → {script_reel}.py {' '.join(extra_args)}" if extra_args else f"  ▶ {module}.py")
    print(f"    Commande: {' '.join(cmd)}")
    
    t0 = time.perf_counter()
    try:
        env = os.environ.copy()
        env['PYTHONIOENCODING'] = 'utf-8'
        
        proc = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='replace',
            timeout=300,
            cwd=str(Path(__file__).parent),
            env=env
        )
        duree = (time.perf_counter() - t0) * 1000
        resultat['duree_ms'] = round(duree)
        
        if proc.returncode == 0:
            resultat['succes'] = True
            if fichier_sortie.exists():
                res = lire_resultats_csv(fichier_sortie)
                resultat['nb_questions'] = len(res)
                print(f"    ✓ {len(res)} question(s) traitée(s) en {resultat['duree_ms']}ms")
                print(f"    → {fichier_sortie.resolve()}")
            else:
                print(f"    ✓ Terminé en {resultat['duree_ms']}ms (fichier sortie non trouvé)")
        else:
            resultat['erreur'] = proc.stderr[:500] if proc.stderr else f"Code retour {proc.returncode}"
            print(f"    ✗ Erreur (code {proc.returncode}) en {resultat['duree_ms']}ms")
            if proc.stderr:
                for line in proc.stderr.strip().split('\n')[:5]:
                    print(f"      {line}")
    
    except subprocess.TimeoutExpired:
        duree = (time.perf_counter() - t0) * 1000
        resultat['duree_ms'] = round(duree)
        resultat['erreur'] = "Timeout (>300s)"
        print(f"    ✗ TIMEOUT après 300s")
    
    except FileNotFoundError:
        resultat['erreur'] = f"Module {script_reel}.py introuvable"
        print(f"    ✗ {script_reel}.py introuvable")
    
    except Exception as e:
        resultat['erreur'] = str(e)[:500]
        print(f"    ✗ Exception: {e}")
    
    return resultat


# ═══════════════════════════════════════════════════════════════════════════
# GÉNÉRATION DU RAPPORT DOCX
# ═══════════════════════════════════════════════════════════════════════════

def generer_rapport_docx(resultats_modules: List[Dict], csv_path: Path,
                         db_path: Optional[Path], questions: List[str],
                         evaluations: List[Dict] = None,
                         verbose: bool = False):
    """
    Génère un rapport DOCX professionnel à partir des résultats.
    
    V2.0 : Inclut la section d'évaluation des commentaires si disponible.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn
    except ImportError:
        print()
        print("⚠ python-docx non installé. Installation...")
        subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx',
                       '--break-system-packages', '-q'], capture_output=True)
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn
    
    from datetime import datetime
    
    doc = Document()
    
    # ─── Style par défaut ───
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    
    # ─── Marges ───
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ═══════════════════════════════════════════════════════════════════
    # PAGE DE TITRE
    # ═══════════════════════════════════════════════════════════════════
    
    for _ in range(4):
        doc.add_paragraph()
    
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titre.add_run("📋 Rapport de Tests")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)
    run.bold = True
    
    sous_titre = doc.add_paragraph()
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sous_titre.add_run("KitView Search — Détection et Recherche")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_paragraph()
    
    infos = doc.add_paragraph()
    infos.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = infos.add_run(f"Fichier de test : {csv_path.name}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    if db_path:
        infos2 = doc.add_paragraph()
        infos2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = infos2.add_run(f"Base de données : {db_path.name}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    resultats_list = resultats_modules if isinstance(resultats_modules, list) else list(resultats_modules.values())
    
    infos3 = doc.add_paragraph()
    infos3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = infos3.add_run(f"{len(questions)} question(s) testée(s) — {len(resultats_list)} module(s)")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # V2.0 : Résumé évaluation sur la page de titre
    if evaluations:
        nb_ok = sum(1 for e in evaluations if e['statut'] == 'OK')
        nb_ko = sum(1 for e in evaluations if e['statut'] == 'KO')
        nb_action = sum(1 for e in evaluations if e['statut'] == 'ACTION')
        nb_pass = sum(1 for e in evaluations if e['statut'] == 'PASS')
        
        infos_eval = doc.add_paragraph()
        infos_eval.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = infos_eval.add_run(f"Évaluation : {nb_ok} OK · {nb_ko} KO · {nb_action} ACTION · {nb_pass} PASS")
        run.font.size = Pt(12)
        if nb_ko > 0:
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        elif nb_action > 0:
            run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
        else:
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        run.bold = True
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(datetime.now().strftime("%d/%m/%Y à %H:%M"))
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # LÉGENDE DES COULEURS
    # ═══════════════════════════════════════════════════════════════════
    
    h = doc.add_heading("🎨  Légende des couleurs", level=1)
    _colorer_heading(h, RGBColor(0x1A, 0x47, 0x7A))
    
    p = doc.add_paragraph()
    run = p.add_run("Les résultats dans les tableaux sont colorés selon le type de critère détecté :")
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    legende = [
        ("[tag]",         "Pathologie / tag médical",          RGBColor(0x8B, 0x00, 0x00)),
        ("[age] / [sexe]","Critère d'âge ou de sexe",          RGBColor(0x00, 0x64, 0x00)),
        ("[angle]",       "Angle céphalométrique (ANB, SNA…)", RGBColor(0x80, 0x00, 0x80)),
        ("[id]",          "Identifiant patient",               RGBColor(0x00, 0x00, 0x8B)),
        ("[meme]",        "Similarité (même X que…)",          RGBColor(0xFF, 0x8C, 0x00)),
        ("[adj]",         "Adjectif qualificatif",             RGBColor(0x69, 0x69, 0x69)),
        ("→ / Résidu",    "Résultat final ou résidu",          RGBColor(0x55, 0x55, 0x55)),
        ("✗ / erreur",    "Erreur ou échec",                   RGBColor(0xCC, 0x00, 0x00)),
    ]
    
    # V2.0 : Légende des statuts d'évaluation
    if evaluations:
        legende.extend([
            ("✅ OK",       "Commentaire résolu par la détection",    RGBColor(0x00, 0x80, 0x00)),
            ("❌ KO",       "Problème persistant malgré la détection", RGBColor(0xCC, 0x00, 0x00)),
            ("🔧 ACTION",   "Vérification manuelle requise",          RGBColor(0xFF, 0x8C, 0x00)),
            ("⬜ PASS",     "Pas de commentaire, succès du module",   RGBColor(0x88, 0x88, 0x88)),
        ])
    
    table_leg = doc.add_table(rows=1, cols=3)
    table_leg.style = 'Table Grid'
    _cellule_entete(table_leg.rows[0].cells[0], "Préfixe")
    _cellule_entete(table_leg.rows[0].cells[1], "Signification")
    _cellule_entete(table_leg.rows[0].cells[2], "Exemple")
    
    for prefixe, signification, couleur in legende:
        row = table_leg.add_row()
        cell_p = row.cells[0]
        cell_p.text = ''
        p_run = cell_p.paragraphs[0].add_run(prefixe)
        p_run.font.size = Pt(9)
        p_run.font.color.rgb = couleur
        p_run.bold = True
        _cellule_donnee(row.cells[1], signification)
        _cellule_donnee(row.cells[2], "")
    
    for row in table_leg.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if not run.bold or run.font.size == Pt(9):
                        run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # ═══════════════════════════════════════════════════════════════════
    # SOMMAIRE
    # ═══════════════════════════════════════════════════════════════════
    
    h = doc.add_heading("Sommaire", level=1)
    _colorer_heading(h, RGBColor(0x1A, 0x47, 0x7A))
    
    for cat_key, cat in CATEGORIES.items():
        modules_cat = [r for r in resultats_list if r['module'] in cat['modules']]
        if not modules_cat:
            continue
        p = doc.add_paragraph()
        run = p.add_run(f"{cat['emoji']}  {cat['titre']}")
        run.bold = True
        run.font.size = Pt(12)
        for r in modules_cat:
            p2 = doc.add_paragraph(style='List Bullet')
            status = "✓" if r['succes'] else "✗"
            p2.add_run(f"{status}  {r['module']}.py — {DESCRIPTIONS_MODULES.get(r['module'], '')}")
    
    # V2.0 : Sommaire évaluation
    if evaluations and any(e['commentaire'] for e in evaluations):
        p = doc.add_paragraph()
        run = p.add_run("📝  Évaluation des commentaires")
        run.bold = True
        run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # QUESTIONS TESTÉES
    # ═══════════════════════════════════════════════════════════════════
    
    h = doc.add_heading("Questions testées", level=1)
    _colorer_heading(h, RGBColor(0x1A, 0x47, 0x7A))
    
    p = doc.add_paragraph()
    p.add_run(f"Le fichier ").font.size = Pt(11)
    run = p.add_run(csv_path.name)
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(f" contient {len(questions)} question(s) :").font.size = Pt(11)
    
    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"Q{i}: ")
        run.bold = True
        p.add_run(q)
        # V2.0 : Ajouter le commentaire s'il existe
        if evaluations and i - 1 < len(evaluations):
            comm = evaluations[i - 1].get('commentaire', '')
            if comm:
                run_c = p.add_run(f"\n    💬 {comm}")
                run_c.italic = True
                run_c.font.size = Pt(9)
                run_c.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # V2.0 : ÉVALUATION DES COMMENTAIRES
    # ═══════════════════════════════════════════════════════════════════
    
    if evaluations and any(e['commentaire'] for e in evaluations):
        h = doc.add_heading("📝  Évaluation des commentaires", level=1)
        _colorer_heading(h, RGBColor(0x1A, 0x47, 0x7A))
        
        p = doc.add_paragraph()
        run = p.add_run(
            "Chaque commentaire du testeur a été évalué par l'IA "
            "contre les résultats du module detall (pipeline complet de détection)."
        )
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        
        # Tableau d'évaluation
        table_eval = doc.add_table(rows=1, cols=4)
        table_eval.style = 'Table Grid'
        _cellule_entete(table_eval.rows[0].cells[0], "Question")
        _cellule_entete(table_eval.rows[0].cells[1], "Commentaire")
        _cellule_entete(table_eval.rows[0].cells[2], "Statut")
        _cellule_entete(table_eval.rows[0].cells[3], "Explication / Action")
        
        for ev in evaluations:
            if not ev.get('commentaire'):
                continue
            
            row = table_eval.add_row()
            
            # Question (tronquée)
            q_text = ev['question']
            if len(q_text) > 45:
                q_text = q_text[:42] + "..."
            _cellule_donnee(row.cells[0], q_text, bold=True)
            
            # Commentaire
            comm = ev['commentaire']
            if len(comm) > 60:
                comm = comm[:57] + "..."
            _cellule_donnee(row.cells[1], comm)
            
            # Statut avec couleur
            statut = ev['statut']
            cell_statut = row.cells[2]
            cell_statut.text = ''
            p_s = cell_statut.paragraphs[0]
            emoji = {'OK': '✅', 'KO': '❌', 'ACTION': '🔧', 'PASS': '⬜'}.get(statut, '❓')
            run_s = p_s.add_run(f"{emoji} {statut}")
            run_s.bold = True
            run_s.font.size = Pt(9)
            couleurs_statut = {
                'OK': RGBColor(0x00, 0x80, 0x00),
                'KO': RGBColor(0xCC, 0x00, 0x00),
                'ACTION': RGBColor(0xFF, 0x8C, 0x00),
                'PASS': RGBColor(0x88, 0x88, 0x88),
            }
            run_s.font.color.rgb = couleurs_statut.get(statut, RGBColor(0x00, 0x00, 0x00))
            
            # Explication + action
            cell_expl = row.cells[3]
            cell_expl.text = ''
            p_e = cell_expl.paragraphs[0]
            expl = ev.get('explication', '')
            run_e = p_e.add_run(expl)
            run_e.font.size = Pt(8)
            if ev.get('action'):
                run_a = p_e.add_run(f"\n📋 {ev['action']}")
                run_a.font.size = Pt(8)
                run_a.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
                run_a.bold = True
        
        # Réduire taille
        for row in table_eval.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.size is None:
                            run.font.size = Pt(8)
        
        # Résumé évaluation
        doc.add_paragraph()
        nb_ok = sum(1 for e in evaluations if e['statut'] == 'OK')
        nb_ko = sum(1 for e in evaluations if e['statut'] == 'KO')
        nb_action = sum(1 for e in evaluations if e['statut'] == 'ACTION')
        nb_pass = sum(1 for e in evaluations if e['statut'] == 'PASS')
        
        p = doc.add_paragraph()
        run = p.add_run(f"Bilan : {nb_ok} OK · {nb_ko} KO · {nb_action} ACTION · {nb_pass} PASS")
        run.bold = True
        run.font.size = Pt(12)
        if nb_ko > 0:
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        elif nb_action > 0:
            run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
        else:
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        
        doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # RÉSULTATS PAR CATÉGORIE (existant, inchangé)
    # ═══════════════════════════════════════════════════════════════════
    
    resultats_par_module = {r['module']: r for r in resultats_list}
    
    for cat_key, cat in CATEGORIES.items():
        modules_cat = [m for m in cat['modules'] if m in resultats_par_module]
        if not modules_cat:
            continue
        
        h = doc.add_heading(f"{cat['emoji']}  {cat['titre']}", level=1)
        _colorer_heading(h, RGBColor(0x1A, 0x47, 0x7A))
        
        p = doc.add_paragraph()
        run = p.add_run(cat['description'])
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        
        for module in modules_cat:
            r = resultats_par_module[module]
            
            h2 = doc.add_heading(f"{module}.py", level=2)
            _colorer_heading(h2, RGBColor(0x2E, 0x75, 0xB6))
            
            desc = DESCRIPTIONS_MODULES.get(module, '')
            if desc:
                p = doc.add_paragraph()
                run = p.add_run(desc)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            
            if r['succes']:
                p = doc.add_paragraph()
                run = p.add_run(f"✅ Succès — {r['nb_questions']} question(s) en {r['duree_ms']}ms")
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                run.bold = True
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"❌ Échec — {r.get('erreur', 'Erreur inconnue')}")
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                run.bold = True
                continue
            
            fichier_sortie = r.get('fichier_sortie')
            if fichier_sortie and Path(fichier_sortie).exists():
                resultats_csv = lire_resultats_csv(Path(fichier_sortie))
                
                if resultats_csv:
                    max_l = max(len(res['lignes']) for res in resultats_csv)
                    nb_cols = 1 + max_l
                    table = doc.add_table(rows=1, cols=nb_cols)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    
                    hdr = table.rows[0]
                    _cellule_entete(hdr.cells[0], "Question")
                    for i in range(max_l):
                        _cellule_entete(hdr.cells[1 + i], f"L{i + 1}")
                    
                    for res in resultats_csv:
                        row = table.add_row()
                        q_text = res['question']
                        if len(q_text) > 50:
                            q_text = q_text[:47] + "..."
                        _cellule_donnee(row.cells[0], q_text, bold=True)
                        
                        for i in range(max_l):
                            texte = res['lignes'][i] if i < len(res['lignes']) else ''
                            cell = row.cells[1 + i]
                            _cellule_donnee_coloree(cell, texte)
                    
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(8)
                    
                    doc.add_paragraph()
        
        # Tableau comparatif pour les recherches
        if cat_key == 'recherche':
            _generer_tableau_comparatif(
                doc, resultats_par_module, questions,
                groupes=[
                    ('trouve', 'standard'),
                    ('trouve_ia', 'IA (detia)'),
                    ('trouve_iabrut', 'IA brut (none)'),
                ],
                titre="Comparatif trouve : standard vs IA vs IA brut"
            )
            _generer_tableau_comparatif(
                doc, resultats_par_module, questions,
                groupes=[
                    ('search', 'standard+fallback'),
                    ('search_purstandard', 'purstandard'),
                    ('search_puria', 'puria'),
                ],
                titre="Comparatif search : modes de routage"
            )
        
        doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════
    # SYNTHÈSE GLOBALE
    # ═══════════════════════════════════════════════════════════════════
    
    h = doc.add_heading("📊  Synthèse globale", level=1)
    _colorer_heading(h, RGBColor(0x1A, 0x47, 0x7A))
    
    nb_ok_mod = sum(1 for r in resultats_list if r['succes'])
    nb_ko_mod = len(resultats_list) - nb_ok_mod
    duree_totale = sum(r['duree_ms'] for r in resultats_list)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    _cellule_entete(table.rows[0].cells[0], "Module")
    _cellule_entete(table.rows[0].cells[1], "Statut")
    _cellule_entete(table.rows[0].cells[2], "Questions")
    _cellule_entete(table.rows[0].cells[3], "Durée")
    
    for r in resultats_list:
        row = table.add_row()
        _cellule_donnee(row.cells[0], r['module'], bold=True)
        if r['succes']:
            _cellule_donnee(row.cells[1], "✅ OK")
        else:
            _cellule_donnee(row.cells[1], "❌ KO")
        _cellule_donnee(row.cells[2], str(r['nb_questions']))
        _cellule_donnee(row.cells[3], f"{r['duree_ms']}ms")
    
    row = table.add_row()
    _cellule_donnee(row.cells[0], "TOTAL", bold=True)
    _cellule_donnee(row.cells[1], f"{nb_ok_mod} ✅  {nb_ko_mod} ❌", bold=True)
    _cellule_donnee(row.cells[2], str(len(questions)), bold=True)
    _cellule_donnee(row.cells[3], f"{duree_totale}ms", bold=True)
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # Conclusion modules
    p = doc.add_paragraph()
    if nb_ko_mod == 0:
        run = p.add_run(f"🎉 Tous les {nb_ok_mod} modules ont réussi leurs tests !")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        run.bold = True
    else:
        run = p.add_run(f"⚠ {nb_ko_mod} module(s) en échec sur {len(resultats_list)}")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0xCC, 0x80, 0x00)
        run.bold = True
    
    # V2.0 : Conclusion évaluation commentaires
    if evaluations and any(e['commentaire'] for e in evaluations):
        doc.add_paragraph()
        nb_ok_e = sum(1 for e in evaluations if e['statut'] == 'OK')
        nb_ko_e = sum(1 for e in evaluations if e['statut'] == 'KO')
        nb_act_e = sum(1 for e in evaluations if e['statut'] == 'ACTION')
        
        p = doc.add_paragraph()
        if nb_ko_e == 0 and nb_act_e == 0:
            run = p.add_run(f"✅ Tous les commentaires évalués sont résolus ({nb_ok_e} OK)")
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        elif nb_ko_e > 0:
            run = p.add_run(f"❌ {nb_ko_e} problème(s) persistant(s) détecté(s) par l'évaluation IA")
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        else:
            run = p.add_run(f"🔧 {nb_act_e} vérification(s) manuelle(s) requise(s)")
            run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
        run.font.size = Pt(12)
        run.bold = True
    
    # ─── Sauvegarde dans tests/ comme {stem}_init.docx ───
    stem = csv_path.stem
    tests_dir = Path(__file__).parent / "tests"
    tests_dir.mkdir(exist_ok=True)
    rapport_path = tests_dir / f"{stem}_init.docx"
    doc.save(str(rapport_path))
    
    return rapport_path


# ═══════════════════════════════════════════════════════════════════════════
# V3.0 : LECTURE DU FICHIER DE BUGS
# ═══════════════════════════════════════════════════════════════════════════

def lire_bugs_docx(bugs_path: Path, debug=False) -> List[Dict]:
    """
    Lit un fichier bugs{stem}.docx et extrait les bugs structurés par test.
    
    Le fichier est structuré : paragraphe "Test N : question" suivi d'une
    table à 2 colonnes (Bugs constatés | Corrections effectuées).
    
    Retourne une liste de dicts :
    [
        {
            'titre': 'Test 1 : patients en classe 2 agés...',
            'bugs': ['detage.py : Le tracking...', 'detia.py : L\'IA fusionnait...'],
            'corrections': ['detage.py : Correction du tracking...', 'Prompt IA : Ajout...'],
        },
        ...
    ]
    """
    try:
        from docx import Document
        from docx.table import Table as DocxTable
    except ImportError:
        print("  ⚠ python-docx non disponible pour lire le fichier de bugs")
        return []
    
    try:
        doc = Document(str(bugs_path))
    except Exception as e:
        print(f"  ⚠ Erreur lecture {bugs_path}: {e}")
        return []
    
    tests = []
    current_test = None
    
    # Itérer sur les éléments du body (paragraphes + tables en séquence)
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        
        if tag == 'p':
            texte = (element.text or '').strip()
            if not texte:
                continue
            
            # Détecter "Test N :" → nouveau test
            if texte.lower().startswith('test ') and ':' in texte[:15]:
                if current_test:
                    tests.append(current_test)
                # Extraire la question (après "Test N : ")
                current_test = {
                    'titre': texte,
                    'bugs': [],
                    'corrections': [],
                }
            # Section post-tests → on arrête la collecte
            elif current_test and any(mot in texte.lower() for mot in
                                       ['correction du prompt', 'disambiguation', 'règle anti-fusion',
                                        'bug search.py']):
                tests.append(current_test)
                current_test = None
        
        elif tag == 'tbl' and current_test is not None:
            # Table associée au test courant : col 0 = bugs, col 1 = corrections
            tbl = DocxTable(element, doc)
            for row_idx, row in enumerate(tbl.rows):
                if row_idx == 0:
                    # Entête ("Bugs constatés" / "Corrections effectuées") → skip
                    first_cell = row.cells[0].text.strip().lower()
                    if 'bug' in first_cell or 'constat' in first_cell:
                        continue
                
                if len(row.cells) >= 2:
                    bug_text = row.cells[0].text.strip()
                    corr_text = row.cells[1].text.strip()
                    if bug_text:
                        current_test['bugs'].append(bug_text)
                    if corr_text:
                        current_test['corrections'].append(corr_text)
    
    if current_test:
        tests.append(current_test)
    
    if debug:
        print(f"[DEBUG] lire_bugs_docx: {len(tests)} test(s) extraits de {bugs_path.name}")
        for t in tests:
            print(f"  {t['titre'][:60]} → {len(t['bugs'])} bug(s), {len(t['corrections'])} correction(s)")
    
    return tests


# ═══════════════════════════════════════════════════════════════════════════
# V3.0 : GÉNÉRATION DU RAPPORT FINAL (~5 PAGES)
# ═══════════════════════════════════════════════════════════════════════════

def generer_rapport_final_docx(resultats_modules: List[Dict], csv_path: Path,
                                db_path: Path, questions: List[str],
                                evaluations: List[Dict] = None,
                                bugs: List[Dict] = None,
                                verbose: bool = False):
    """
    Génère {stem}_final.docx dans tests/ — rapport compact ~5 pages.
    
    Partie 1 : Résumé des bugs corrigés (depuis bugs{stem}.docx)
    Partie 2 : Résumé des résultats après corrections
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from datetime import datetime
    
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    resultats_list = resultats_modules if isinstance(resultats_modules, list) else list(resultats_modules.values())
    nb_ok_mod = sum(1 for r in resultats_list if r['succes'])
    nb_ko_mod = len(resultats_list) - nb_ok_mod
    duree_totale = sum(r['duree_ms'] for r in resultats_list)
    
    # ═══════════════════════════════════════════════════════════════════
    # EN-TÊTE COMPACT (pas de page de titre pleine page)
    # ═══════════════════════════════════════════════════════════════════
    
    doc.add_paragraph()
    
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titre.add_run("📋 Rapport Final — Après corrections")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)
    run.bold = True
    
    sous_titre = doc.add_paragraph()
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sous_titre.add_run(f"KitView Search — {csv_path.name} / {db_path.name if db_path else '?'}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    infos = doc.add_paragraph()
    infos.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = infos.add_run(f"{len(questions)} questions · {len(resultats_list)} modules · {nb_ok_mod} ✅ {nb_ko_mod} ❌")
    run.font.size = Pt(11)
    
    if evaluations:
        nb_eval_ok = sum(1 for e in evaluations if e['statut'] == 'OK')
        nb_eval_ko = sum(1 for e in evaluations if e['statut'] == 'KO')
        nb_eval_action = sum(1 for e in evaluations if e['statut'] == 'ACTION')
        infos_eval = doc.add_paragraph()
        infos_eval.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = infos_eval.add_run(f"Évaluation commentaires : {nb_eval_ok} OK · {nb_eval_ko} KO · {nb_eval_action} ACTION")
        run.font.size = Pt(11)
        if nb_eval_ko == 0:
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        else:
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run.bold = True
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(datetime.now().strftime("%d/%m/%Y à %H:%M"))
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True
    
    doc.add_paragraph()
    
    # ═══════════════════════════════════════════════════════════════════
    # PARTIE 1 : BUGS CORRIGÉS
    # ═══════════════════════════════════════════════════════════════════
    
    h = doc.add_heading("🔧  Partie 1 — Bugs corrigés", level=1)
    _colorer_heading(h, RGBColor(0x1A, 0x47, 0x7A))
    
    if bugs:
        p = doc.add_paragraph()
        run = p.add_run(f"{len(bugs)} test(s) ont révélé des bugs, tous corrigés. "
                        f"Détails dans le rapport _init et dans bugs{csv_path.stem}.docx.")
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.size = Pt(10)
        
        # Tableau résumé des bugs
        table_bugs = doc.add_table(rows=1, cols=3)
        table_bugs.style = 'Table Grid'
        _cellule_entete(table_bugs.rows[0].cells[0], "Question")
        _cellule_entete(table_bugs.rows[0].cells[1], "Bug(s)")
        _cellule_entete(table_bugs.rows[0].cells[2], "Correction(s)")
        
        for test in bugs:
            row = table_bugs.add_row()
            
            # Extraire la question du titre ("Test N : question...")
            titre_test = test['titre']
            if ':' in titre_test:
                q_part = titre_test.split(':', 1)[1].strip()
            else:
                q_part = titre_test
            if q_part.startswith('*'):
                q_part = q_part.strip('*').strip()
            if len(q_part) > 50:
                q_part = q_part[:47] + "..."
            _cellule_donnee(row.cells[0], q_part, bold=True)
            
            # Bugs résumés
            bugs_txt = []
            for b in test['bugs'][:3]:
                # Extraire le module et le résumé court
                if ':' in b[:30]:
                    module_part = b.split(':')[0].strip().rstrip('*').lstrip('*')
                    desc = b.split(':', 1)[1].strip()[:80]
                    bugs_txt.append(f"• {module_part}: {desc}")
                else:
                    bugs_txt.append(f"• {b[:90]}")
            _cellule_donnee(row.cells[1], '\n'.join(bugs_txt) if bugs_txt else '(aucun)')
            
            # Corrections résumées
            corr_txt = []
            for c in test['corrections'][:3]:
                if ':' in c[:30]:
                    module_part = c.split(':')[0].strip().rstrip('*').lstrip('*')
                    desc = c.split(':', 1)[1].strip()[:80]
                    corr_txt.append(f"• {module_part}: {desc}")
                else:
                    corr_txt.append(f"• {c[:90]}")
            _cellule_donnee(row.cells[2], '\n'.join(corr_txt) if corr_txt else '(aucune)')
        
        for row in table_bugs.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.size is None or run.font.size > Pt(8):
                            run.font.size = Pt(8)
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"⚠ Fichier bugs{csv_path.stem}.docx non trouvé. "
                        "Section renseignée manuellement si nécessaire.")
        run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
        run.italic = True
    
    doc.add_paragraph()
    
    # ═══════════════════════════════════════════════════════════════════
    # PARTIE 2 : RÉSULTATS APRÈS CORRECTIONS
    # ═══════════════════════════════════════════════════════════════════
    
    h = doc.add_heading("✅  Partie 2 — Résultats après corrections", level=1)
    _colorer_heading(h, RGBColor(0x1A, 0x47, 0x7A))
    
    # ─── Évaluation des commentaires ───
    if evaluations and any(e['commentaire'] for e in evaluations):
        h2 = doc.add_heading("Évaluation des commentaires", level=2)
        _colorer_heading(h2, RGBColor(0x2E, 0x75, 0xB6))
        
        table_eval = doc.add_table(rows=1, cols=4)
        table_eval.style = 'Table Grid'
        _cellule_entete(table_eval.rows[0].cells[0], "Question")
        _cellule_entete(table_eval.rows[0].cells[1], "Commentaire")
        _cellule_entete(table_eval.rows[0].cells[2], "Statut")
        _cellule_entete(table_eval.rows[0].cells[3], "Explication")
        
        for ev in evaluations:
            if not ev.get('commentaire'):
                continue
            row = table_eval.add_row()
            q_text = ev['question'][:42] + "..." if len(ev['question']) > 45 else ev['question']
            _cellule_donnee(row.cells[0], q_text, bold=True)
            
            comm = ev['commentaire'][:57] + "..." if len(ev['commentaire']) > 60 else ev['commentaire']
            _cellule_donnee(row.cells[1], comm)
            
            statut = ev['statut']
            cell_s = row.cells[2]
            cell_s.text = ''
            emoji = {'OK': '✅', 'KO': '❌', 'ACTION': '🔧', 'PASS': '⬜'}.get(statut, '❓')
            run_s = cell_s.paragraphs[0].add_run(f"{emoji} {statut}")
            run_s.bold = True
            run_s.font.size = Pt(9)
            couleurs = {'OK': RGBColor(0x00, 0x80, 0x00), 'KO': RGBColor(0xCC, 0x00, 0x00),
                        'ACTION': RGBColor(0xFF, 0x8C, 0x00), 'PASS': RGBColor(0x88, 0x88, 0x88)}
            run_s.font.color.rgb = couleurs.get(statut, RGBColor(0, 0, 0))
            
            _cellule_donnee(row.cells[3], ev.get('explication', '')[:100])
        
        for row in table_eval.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.size is None:
                            run.font.size = Pt(8)
        
        doc.add_paragraph()
    
    # ─── Résultats clés : detall (pipeline complet) ───
    resultats_par_module = {r['module']: r for r in resultats_list}
    
    modules_cles = ['detall', 'detia']
    for module in modules_cles:
        r = resultats_par_module.get(module)
        if not r or not r.get('succes') or not r.get('fichier_sortie'):
            continue
        
        fichier = Path(r['fichier_sortie'])
        if not fichier.exists():
            continue
        
        resultats_csv = lire_resultats_csv(fichier)
        if not resultats_csv:
            continue
        
        h2 = doc.add_heading(f"Résultats {module}.py", level=2)
        _colorer_heading(h2, RGBColor(0x2E, 0x75, 0xB6))
        
        p = doc.add_paragraph()
        run = p.add_run(f"✅ {r['nb_questions']} question(s) en {r['duree_ms']}ms — {DESCRIPTIONS_MODULES.get(module, '')}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        
        # Filtrer les lignes JSON/SQL
        max_l = max(len([l for l in res['lignes'] if not l.startswith('JSON:') and not l.startswith('SQL:')]) for res in resultats_csv)
        nb_cols = 1 + max_l
        table = doc.add_table(rows=1, cols=nb_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        _cellule_entete(table.rows[0].cells[0], "Question")
        for ci in range(max_l):
            _cellule_entete(table.rows[0].cells[1 + ci], f"L{ci + 1}")
        
        for res in resultats_csv:
            row = table.add_row()
            q_text = res['question'][:47] + "..." if len(res['question']) > 50 else res['question']
            _cellule_donnee(row.cells[0], q_text, bold=True)
            
            lignes_filtrees = [l for l in res['lignes'] if not l.startswith('JSON:') and not l.startswith('SQL:')]
            for ci in range(max_l):
                texte = lignes_filtrees[ci] if ci < len(lignes_filtrees) else ''
                _cellule_donnee_coloree(row.cells[1 + ci], texte)
        
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
        
        doc.add_paragraph()
    
    # ─── Synthèse globale compacte ───
    h2 = doc.add_heading("Synthèse des modules", level=2)
    _colorer_heading(h2, RGBColor(0x2E, 0x75, 0xB6))
    
    table_synth = doc.add_table(rows=1, cols=4)
    table_synth.style = 'Table Grid'
    _cellule_entete(table_synth.rows[0].cells[0], "Module")
    _cellule_entete(table_synth.rows[0].cells[1], "Statut")
    _cellule_entete(table_synth.rows[0].cells[2], "Questions")
    _cellule_entete(table_synth.rows[0].cells[3], "Durée")
    
    for r in resultats_list:
        row = table_synth.add_row()
        _cellule_donnee(row.cells[0], r['module'], bold=True)
        _cellule_donnee(row.cells[1], "✅ OK" if r['succes'] else "❌ KO")
        _cellule_donnee(row.cells[2], str(r['nb_questions']))
        _cellule_donnee(row.cells[3], f"{r['duree_ms']}ms")
    
    row_total = table_synth.add_row()
    _cellule_donnee(row_total.cells[0], "TOTAL", bold=True)
    _cellule_donnee(row_total.cells[1], f"{nb_ok_mod} ✅  {nb_ko_mod} ❌", bold=True)
    _cellule_donnee(row_total.cells[2], str(len(questions)), bold=True)
    _cellule_donnee(row_total.cells[3], f"{duree_totale}ms", bold=True)
    
    for row in table_synth.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # ─── Conclusion ───
    p = doc.add_paragraph()
    if nb_ko_mod == 0:
        run = p.add_run(f"🎉 Tous les {nb_ok_mod} modules ont réussi !")
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    else:
        run = p.add_run(f"⚠ {nb_ko_mod} module(s) en échec sur {len(resultats_list)}")
        run.font.color.rgb = RGBColor(0xCC, 0x80, 0x00)
    run.font.size = Pt(13)
    run.bold = True
    
    if evaluations and any(e['commentaire'] for e in evaluations):
        nb_ok_e = sum(1 for e in evaluations if e['statut'] == 'OK')
        nb_ko_e = sum(1 for e in evaluations if e['statut'] == 'KO')
        p2 = doc.add_paragraph()
        if nb_ko_e == 0:
            run2 = p2.add_run(f"✅ Tous les commentaires évalués sont résolus ({nb_ok_e} OK)")
            run2.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        else:
            run2 = p2.add_run(f"❌ {nb_ko_e} problème(s) persistant(s)")
            run2.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run2.font.size = Pt(12)
        run2.bold = True
    
    # ─── Sauvegarde dans tests/ ───
    stem = csv_path.stem
    tests_dir = Path(__file__).parent / "tests"
    tests_dir.mkdir(exist_ok=True)
    rapport_path = tests_dir / f"{stem}_final.docx"
    doc.save(str(rapport_path))
    
    return rapport_path




def _colorer_heading(heading, color: 'RGBColor'):
    """Colore tous les runs d'un heading."""
    for run in heading.runs:
        run.font.color.rgb = color


def _cellule_entete(cell, texte):
    """Formate une cellule d'entête de tableau."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texte)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): '1A477A',
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def _cellule_donnee(cell, texte, bold=False):
    """Formate une cellule de données."""
    from docx.shared import Pt
    
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(texte)
    run.font.size = Pt(9)
    run.bold = bold


def _cellule_donnee_coloree(cell, texte):
    """Formate une cellule avec coloration selon le contenu."""
    from docx.shared import Pt, RGBColor
    
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(texte)
    run.font.size = Pt(8)
    
    texte_lower = texte.lower()
    if texte.startswith('[tag]'):
        run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    elif texte.startswith('[age]') or texte.startswith('[sexe]'):
        run.font.color.rgb = RGBColor(0x00, 0x64, 0x00)
    elif texte.startswith('[angle]'):
        run.font.color.rgb = RGBColor(0x80, 0x00, 0x80)
    elif texte.startswith('[id]'):
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x8B)
    elif texte.startswith('[meme]'):
        run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
    elif texte.startswith('[adj]'):
        run.font.color.rgb = RGBColor(0x69, 0x69, 0x69)
    elif texte.startswith('→') or texte.startswith('Résidu'):
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.italic = True
    elif '✗' in texte or 'erreur' in texte_lower:
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    elif texte.startswith('Routage:'):
        run.font.color.rgb = RGBColor(0x00, 0x64, 0x00)
        run.bold = True


def _extraire_nb_patients(lignes: List[str]) -> Optional[int]:
    """Extrait le nombre de patients d'une ligne '→ N patient(s) en Xms'."""
    import re
    for ligne in lignes:
        m = re.match(r'→\s*(\d+)\s*patient', ligne)
        if m:
            return int(m.group(1))
    return None


def _extraire_routage(lignes: List[str]) -> str:
    """Extrait l'indicateur de routage d'une ligne 'Routage: ...'."""
    for ligne in lignes:
        if ligne.startswith('Routage:'):
            return ligne.replace('Routage:', '').strip()
    return ''


def _generer_tableau_comparatif(doc, resultats_par_module: Dict, questions: List[str],
                                 groupes: List[Tuple[str, str]], titre: str):
    """Génère un tableau comparatif entre plusieurs variantes d'un même module."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    
    modules_avec_resultats = []
    for module_name, label in groupes:
        r = resultats_par_module.get(module_name)
        if r and r.get('succes') and r.get('fichier_sortie') and Path(r['fichier_sortie']).exists():
            modules_avec_resultats.append((module_name, label))
    
    if len(modules_avec_resultats) < 2:
        return
    
    h3 = doc.add_heading(f"📊 {titre}", level=3)
    _colorer_heading(h3, RGBColor(0x2E, 0x75, 0xB6))
    
    p = doc.add_paragraph()
    run = p.add_run("Comparaison du nombre de patients trouvés par chaque méthode :")
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    resultats_par_variante = {}
    for module_name, label in modules_avec_resultats:
        r = resultats_par_module[module_name]
        res_csv = lire_resultats_csv(Path(r['fichier_sortie']))
        resultats_par_variante[module_name] = res_csv
    
    nb_cols = 1 + len(modules_avec_resultats)
    table = doc.add_table(rows=1, cols=nb_cols)
    table.style = 'Table Grid'
    
    _cellule_entete(table.rows[0].cells[0], "Question")
    for i, (_, label) in enumerate(modules_avec_resultats):
        _cellule_entete(table.rows[0].cells[1 + i], label)
    
    for q_idx, question in enumerate(questions):
        row = table.add_row()
        q_text = question[:45] + "..." if len(question) > 45 else question
        _cellule_donnee(row.cells[0], q_text, bold=True)
        
        for i, (module_name, label) in enumerate(modules_avec_resultats):
            res_list = resultats_par_variante.get(module_name, [])
            if q_idx < len(res_list):
                lignes = res_list[q_idx].get('lignes', [])
                nb = _extraire_nb_patients(lignes)
                routage = _extraire_routage(lignes)
                
                cell = row.cells[1 + i]
                cell.text = ''
                p_cell = cell.paragraphs[0]
                
                if nb is not None:
                    texte = str(nb)
                    if routage:
                        texte += f" {routage}"
                    run = p_cell.add_run(texte)
                    run.font.size = Pt(9)
                    if nb > 0:
                        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                        run.bold = True
                    else:
                        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                else:
                    run = p_cell.add_run("—")
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            else:
                _cellule_donnee(row.cells[1 + i], "—")
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
# AIDE CLI
# ═══════════════════════════════════════════════════════════════════════════

def afficher_aide():
    """Affiche l'aide."""
    print("Usage:")
    print(f"  python {__pgm__} <fichier.csv> <base.db> [options]")
    print(f"  python {__pgm__} <base.db> <fichier.csv> [options]")
    print()
    print("Arguments (ordre libre, TOUS DEUX OBLIGATOIRES) :")
    print("  *.csv     Fichier de test (cherche aussi dans tests/)")
    print("  *.db      Base SQLite (cherche aussi dans bases/)")
    print()
    print("Options :")
    print("  -v         Mode verbose")
    print("  -d         Mode debug")
    print("  --no-ia    Exclure les modules IA (detia, detiabrut)")
    print("  --no-db    Exclure les modules nécessitant une base")
    print("  --no-doc   Pas de génération du rapport DOCX")
    print("  --no-eval  Pas d'évaluation IA des commentaires")
    print("  --only=x   Exécuter uniquement le module x")
    print()
    print("Rapports DOCX (dans tests/) :")
    print("─" * 60)
    print("  1er lancement → tests/{stem}_init.docx  (rapport complet)")
    print("  2e  lancement → tests/{stem}_final.docx (résumé ~5 pages)")
    print("  Le rapport final utilise tests/bugs{stem}.docx si présent")
    print()
    print("Modules disponibles :")
    print("─" * 60)
    print("  Détection standard (sans base) :")
    for m in ['detcount', 'detid', 'detage', 'detangles', 'dettags', 'detmeme', 'detall']:
        print(f"    {m:20} {DESCRIPTIONS_MODULES.get(m, '')}")
    print()
    print("  Détection IA (sans base) :")
    for m in ['detia', 'detiabrut']:
        print(f"    {m:20} {DESCRIPTIONS_MODULES.get(m, '')}")
    print()
    print("  Recherche (avec base) :")
    for m in ['trouve', 'trouve_ia', 'trouve_iabrut', 'trouveid',
              'search', 'search_purstandard', 'search_puria']:
        print(f"    {m:20} {DESCRIPTIONS_MODULES.get(m, '')}")
    print()
    print("Évaluation des commentaires :")
    print("─" * 60)
    print("  Si le CSV contient une colonne 'commentaire', chaque commentaire")
    print("  est évalué par l'IA contre les résultats de detall.")
    print("  Statuts : OK (résolu) / KO (problème persistant) / ACTION (vérif manuelle)")
    print("  Désactiver : --no-eval")
    print()
    print("Exemples :")
    print("─" * 60)
    print(f"  python {__pgm__} quentin.csv base1964.db          # → quentin_init.docx")
    print(f"  python {__pgm__} quentin.csv base1964.db          # → quentin_final.docx (si _init existe)")
    print(f"  python {__pgm__} quentin.csv base1964.db --no-ia")
    print(f"  python {__pgm__} quentin.csv base1964.db --only=detage")
    print()


# ═══════════════════════════════════════════════════════════════════════════
# POINT D'ENTRÉE
# ═══════════════════════════════════════════════════════════════════════════

def main():
    """Point d'entrée CLI."""
    print(f"╔════════════════════════════════════════════════════════════════")
    print(f"║ {__pgm__} V{__version__} - {__date__}")
    print(f"║ Exécution complète de tous les modules de test")
    print(f"║ V3.0 : Rapports _init / _final dans tests/")
    print(f"╚════════════════════════════════════════════════════════════════")
    print()
    
    if len(sys.argv) < 2:
        afficher_aide()
        return 1
    
    # ─── Parser les arguments ───
    csv_arg = None
    db_arg = None
    verbose = False
    debug = False
    no_ia = False
    no_db = False
    no_doc = False
    no_eval = False
    only_module = None
    
    for arg in sys.argv[1:]:
        if arg in ('-v', '--verbose'):
            verbose = True
        elif arg in ('-d', '--debug'):
            debug = True
        elif arg == '--no-ia':
            no_ia = True
        elif arg == '--no-db':
            no_db = True
        elif arg == '--no-doc':
            no_doc = True
        elif arg == '--no-eval':
            no_eval = True
        elif arg.startswith('--only='):
            only_module = arg.split('=', 1)[1].strip()
        elif arg.endswith('.db'):
            db_arg = arg
        elif arg.endswith('.csv'):
            csv_arg = arg
        elif not arg.startswith('-'):
            csv_arg = arg
    
    # ─── Vérifications ───
    if not csv_arg or not db_arg:
        print("ERREUR : Les deux arguments sont obligatoires (CSV + base)")
        print(f"  Exemple: python {__pgm__} quentin.csv base1964.db")
        return 1
    
    csv_path = _trouver_csv(csv_arg)
    if not csv_path:
        print(f"ERREUR : Fichier {csv_arg} introuvable")
        print(f"  Chemins testés:")
        print(f"    - {os.path.abspath(csv_arg)}")
        script_dir = Path(__file__).parent
        print(f"    - {(script_dir / 'tests' / Path(csv_arg).name).resolve()}")
        return 1
    
    db_path = _trouver_db(db_arg)
    if not db_path:
        print(f"ERREUR : Base {db_arg} introuvable")
        print(f"  Chemins testés:")
        print(f"    - {os.path.abspath(db_arg)}")
        script_dir = Path(__file__).parent
        print(f"    - {(script_dir / 'bases' / Path(db_arg).name).resolve()}")
        return 1
    
    # ─── Charger la configuration ───
    print("Configuration :")
    modules = charger_modules_config(verbose=True, debug=debug)
    
    if only_module:
        if only_module in modules:
            modules = [only_module]
        else:
            print(f"ERREUR : Module '{only_module}' non trouvé dans la liste")
            print(f"  Modules disponibles: {', '.join(sorted(modules))}")
            return 1
    
    if no_ia:
        modules = [m for m in modules if m not in MODULES_IA]
        print(f"  --no-ia : modules IA exclus")
    
    if no_db:
        modules = [m for m in modules if m not in MODULES_DB]
        print(f"  --no-db : modules base exclus")
    elif not db_path:
        modules_exclus = [m for m in modules if m in MODULES_DB]
        modules = [m for m in modules if m not in MODULES_DB]
        if modules_exclus:
            print(f"  Pas de base .db → exclusion auto: {', '.join(modules_exclus)}")
    
    modules.sort()
    
    print()
    print(f"Fichier CSV  : {csv_path}")
    if db_path:
        print(f"Base         : {db_path}")
    print(f"Modules ({len(modules)}) : {', '.join(modules)}")
    
    # ─── Lire les questions ET commentaires ───
    questions = lire_questions_csv(csv_path)
    commentaires = lire_commentaires_csv(csv_path)
    print(f"Questions    : {len(questions)}")
    
    nb_commentaires = sum(1 for c in commentaires if c) if commentaires else 0
    if commentaires and nb_commentaires > 0:
        print(f"Commentaires : {nb_commentaires} (évaluation IA {'désactivée' if no_eval else 'activée'})")
    
    print()
    
    if not questions:
        print("ERREUR : Aucune question trouvée dans le fichier CSV")
        return 1
    
    # ═══════════════════════════════════════════════════════════════════
    # EXÉCUTION DES MODULES
    # ═══════════════════════════════════════════════════════════════════
    
    print("═" * 64)
    print(f"  EXÉCUTION DE {len(modules)} MODULE(S)")
    print("═" * 64)
    
    t_global = time.perf_counter()
    resultats = []
    
    for module in modules:
        r = executer_module(module, csv_path, db_path, verbose=verbose, debug=debug)
        resultats.append(r)
    
    duree_globale = round((time.perf_counter() - t_global) * 1000)
    
    # ═══════════════════════════════════════════════════════════════════
    # ÉVALUATION DES COMMENTAIRES (V2.0)
    # ═══════════════════════════════════════════════════════════════════
    
    evaluations = None
    
    if commentaires and nb_commentaires > 0 and not no_eval:
        print()
        print("═" * 64)
        print(f"  ÉVALUATION IA DES COMMENTAIRES ({nb_commentaires} commentaire(s))")
        print("═" * 64)
        print()
        
        # Construire le dict des résultats par module
        resultats_par_module = {r['module']: r for r in resultats}
        
        evaluations = evaluer_tous_commentaires(
            questions, commentaires, resultats_par_module,
            verbose=verbose, debug=debug
        )
    
    # ═══════════════════════════════════════════════════════════════════
    # RÉSUMÉ CONSOLE
    # ═══════════════════════════════════════════════════════════════════
    
    print()
    print("═" * 64)
    print("  RÉSUMÉ MODULES")
    print("═" * 64)
    
    nb_ok = sum(1 for r in resultats if r['succes'])
    nb_ko = len(resultats) - nb_ok
    
    for r in resultats:
        status = "✓" if r['succes'] else "✗"
        duree = f"{r['duree_ms']}ms"
        print(f"  {status} {r['module']:14} {r['nb_questions']:3} question(s)  {duree:>8}")
    
    print(f"  {'─' * 58}")
    print(f"  {nb_ok} succès, {nb_ko} échec(s) — Durée totale: {duree_globale}ms")
    
    # V2.0 : Résumé évaluation
    if evaluations:
        print()
        print("═" * 64)
        print("  RÉSUMÉ ÉVALUATION COMMENTAIRES")
        print("═" * 64)
        
        nb_eval_ok = sum(1 for e in evaluations if e['statut'] == 'OK')
        nb_eval_ko = sum(1 for e in evaluations if e['statut'] == 'KO')
        nb_eval_action = sum(1 for e in evaluations if e['statut'] == 'ACTION')
        nb_eval_pass = sum(1 for e in evaluations if e['statut'] == 'PASS')
        
        for e in evaluations:
            if not e.get('commentaire'):
                continue
            emoji = {'OK': '✅', 'KO': '❌', 'ACTION': '🔧'}.get(e['statut'], '❓')
            print(f"  {emoji} {e['statut']:6} \"{e['question'][:50]}\"")
            print(f"         {e.get('explication', '')[:70]}")
            if e.get('action'):
                print(f"         📋 {e['action'][:70]}")
        
        print(f"  {'─' * 58}")
        print(f"  {nb_eval_ok} OK, {nb_eval_ko} KO, {nb_eval_action} ACTION, {nb_eval_pass} PASS")
    
    # ═══════════════════════════════════════════════════════════════════
    # V3.0 : GÉNÉRATION DU RAPPORT DOCX (_init ou _final)
    # ═══════════════════════════════════════════════════════════════════
    
    if not no_doc:
        print()
        
        stem = csv_path.stem
        tests_dir = Path(__file__).parent / "tests"
        tests_dir.mkdir(exist_ok=True)
        init_path = tests_dir / f"{stem}_init.docx"
        
        if not init_path.exists():
            # ─── MODE INIT : premier passage, générer le rapport initial ───
            print(f"📋 {stem}_init.docx absent → génération du rapport initial")
            try:
                rapport_path = generer_rapport_docx(
                    resultats, csv_path, db_path, questions,
                    evaluations=evaluations,
                    verbose=verbose
                )
                print(f"  ✓ Rapport initial: {rapport_path}")
            except Exception as e:
                print(f"  ✗ Erreur génération rapport initial: {e}")
                if debug:
                    import traceback
                    traceback.print_exc()
        else:
            # ─── MODE FINAL : _init existe, générer le rapport final ───
            print(f"📋 {stem}_init.docx trouvé → génération du rapport final")
            
            # Chercher bugs{stem}.docx
            bugs_path = tests_dir / f"bugs{stem}.docx"
            bugs = None
            if bugs_path.exists():
                print(f"  📄 Fichier de bugs: {bugs_path}")
                bugs = lire_bugs_docx(bugs_path, debug=debug)
                print(f"     {len(bugs)} test(s) avec bugs extraits")
            else:
                print(f"  ⚠ {bugs_path.name} non trouvé — partie bugs sera vide")
            
            try:
                rapport_path = generer_rapport_final_docx(
                    resultats, csv_path, db_path, questions,
                    evaluations=evaluations,
                    bugs=bugs,
                    verbose=verbose
                )
                print(f"  ✓ Rapport final: {rapport_path}")
            except Exception as e:
                print(f"  ✗ Erreur génération rapport final: {e}")
                if debug:
                    import traceback
                    traceback.print_exc()
    
    # ─── Fichiers générés ───
    print()
    print("Fichiers générés :")
    for r in resultats:
        if r['succes'] and r['fichier_sortie'] and Path(r['fichier_sortie']).exists():
            print(f"  → {r['fichier_sortie']}")
    
    if not no_doc:
        stem = csv_path.stem
        tests_dir = Path(__file__).parent / "tests"
        for suffix in ['_init.docx', '_final.docx']:
            rapport = tests_dir / f"{stem}{suffix}"
            if rapport.exists():
                print(f"  → {rapport}")
    
    print()
    return 0 if nb_ko == 0 else 1


if __name__ == '__main__':
    sys.exit(main())
