#*TO*#
__pgm__ = "quentin_goback.py"
__version__ = "0.0.0"
__date__ = "01/01/1970 00:00"

"""
Génère quentin_init.docx : rapport de tests tel qu'il aurait été AVANT corrections.

Ce script reconstitue ce qu'aurait donné l'exécution de detfull.py sur les 5 questions
de test de Quentin AVANT les corrections de bugs documentées dans bugsquentin.docx.

Bugs reconstitués :
  - detage.py : tracking en caractères au lieu de mots → 1 seul critère d'âge sur 2
  - ages.csv  : patterns manquants ("compris entre", "supérieur à", "inférieur à", "e"/"&")
  - detia.py  : fusion de 2 critères d'âge en BETWEEN, "classe 2" → squelettique
  - search.py : parser batch cassé, langue "zu" → module en échec

Format : identique à quentin_rapport.docx SANS les colonnes JSON/SQL.

SYNTAXE CLI :
    python quentin_goback.py              → Affiche l'aide
    python quentin_goback.py all          → Génère quentin_init.docx
    python quentin_goback.py all -v       → Mode verbose
    python quentin_goback.py all -d       → Mode debug

OPTIONS :
    -v       Mode verbose
    -d       Mode debug
"""

import sys
import os
import subprocess
from pathlib import Path
from datetime import datetime

# ═══════════════════════════════════════════════════════════════════════════
# DONNÉES RECONSTITUÉES (AVANT CORRECTIONS)
# ═══════════════════════════════════════════════════════════════════════════

QUESTIONS = [
    {
        'question': "patients en classe 2 agés de plus de 20 ans et de moins de 23 ans",
        'resultat': "trop de patients",
        'commentaire': "Pas de détection de plus de 20 ans",
    },
    {
        'question': "Patient compris entre 13 et 16 avec une classe d'angle de 1 uniquement",
        'resultat': "aucun filtre",
        'commentaire': "Il ne voit même pas que 16 et 13 sont des ages",
    },
    {
        'question': "Patients avec un age supérieur à 17 ans et inférieur à 23 ans de classe 3",
        'resultat': "312 patients au lieu de 12",
        'commentaire': "là c'est la classe 3 qui a l'air de poser problème",
    },
    {
        'question': "Patient enfant avec rétrognathie et ou classe 2",
        'resultat': "tous les enfants retournés",
        'commentaire': "rétrognatie non prise en compte",
    },
    {
        'question': "Patient metal compris entre 20 e 24\u00a0ans",
        'resultat': "pas d'âge pris en compte",
        'commentaire': "le e à la place du et a tout foiré.",
    },
]

# ─── Résultats par module AVANT corrections ───

# detcount : inchangé (LIST/COUNT fonctionne)
DETCOUNT = [
    ['→ LIST', "Résidu: 'patients en classe 2 ages de plus de 20 ans et de moins de 23 ans'"],
    ['→ LIST', "Résidu: 'patient compris entre 13 et 16 avec une classe d angle de 1 uniquement'"],
    ['→ LIST', "Résidu: 'patients avec un age superieur a 17 ans et inferieur a 23 ans de classe 3'"],
    ['→ LIST', "Résidu: 'patient enfant avec retrognathie et ou classe 2'"],
    ['→ LIST', "Résidu: 'patient metal compris entre 20 e 24 ans'"],
]

# detid : inchangé (pas d'identifiants dans les questions)
DETID = [
    ["Résidu: 'patients en classe 2 ages de plus de 20 ans et de moins de 23 ans'"],
    ["Résidu: 'patient compris entre 13 et 16 avec une classe d angle de 1 uniquement'"],
    ["Résidu: 'patients avec un age superieur a 17 ans et inferieur a 23 ans de classe 3'"],
    ["Résidu: 'patient enfant avec retrognathie et ou classe 2'"],
    ["Résidu: 'patient metal compris entre 20 e 24 ans'"],
]

# detage AVANT : tracking bug (chars au lieu de mots) + patterns manquants
DETAGE = [
    # Q1: Bug tracking → seul "moins de 23 ans" détecté, "plus de 20 ans" manqué
    ['[age] moins de 23 ans', "Résidu: 'patients en classe 2 ages de plus de 20 ans et de'"],
    # Q2: "compris entre X et Y" absent de ages.csv → aucune détection
    ["Résidu: 'patient compris entre 13 et 16 avec une classe d angle de 1 uniquement'"],
    # Q3: "supérieur à" et "inférieur à" absents de ages.csv → aucune détection
    ["Résidu: 'patients avec un age superieur a 17 ans et inferieur a 23 ans de classe 3'"],
    # Q4: "enfant" fonctionne (pas un pattern ages.csv)
    ['[age] enfant', "Résidu: 'patient avec retrognathie et ou classe 2'"],
    # Q5: "e" au lieu de "et" non toléré → aucune détection
    ["Résidu: 'patient metal compris entre 20 e 24 ans'"],
]

# detangles : inchangé (pas de références ANB/SNA/SNB directes dans les questions)
DETANGLES = [
    ["Résidu: 'patients en classe 2 ages de plus de 20 ans et de moins de 23 ans'"],
    ["Résidu: 'patient compris entre 13 et 16 avec une classe d angle de 1 uniquement'"],
    ["Résidu: 'patients avec un age superieur a 17 ans et inferieur a 23 ans de classe 3'"],
    ["Résidu: 'patient enfant avec retrognathie et ou classe 2'"],
    ["Résidu: 'patient metal compris entre 20 e 24 ans'"],
]

# dettags : fonctionne correctement (le bug était dans detia, pas dans dettags)
DETTAGS = [
    ["[tag] classe ii d'angle (pathologie)", "Résidu: 'patients en ages de plus de 20 ans et de moins de 23 ans'"],
    ["[tag] classe i d'angle (pathologie)", "Résidu: 'patient compris entre 13 et 16 avec une uniquement'"],
    ["[tag] classe iii d'angle (pathologie)", "Résidu: 'patients avec un age superieur a 17 ans et inferieur a 23 ans de'"],
    ["[tag] classe ii d'angle (pathologie)", "[tag] rétrognathie mandibulaire (pathologie)", "Résidu: 'patient enfant avec et ou'"],
    ["[tag] métal (matériau)", "Résidu: 'patient compris entre 20 e 24 ans'"],
]

# detmeme : inchangé (pas de "même X que..." dans les questions)
DETMEME = [
    ["Résidu: 'patients en classe 2 ages de plus de 20 ans et de moins de 23 ans'"],
    ["Résidu: 'patient compris entre 13 et 16 avec une classe d angle de 1 uniquement'"],
    ["Résidu: 'patients avec un age superieur a 17 ans et inferieur a 23 ans de classe 3'"],
    ["Résidu: 'patient enfant avec retrognathie et ou classe 2'"],
    ["Résidu: 'patient metal compris entre 20 e 24 ans'"],
]

# detall AVANT : combine dettags (OK) + detage (buggé)
DETALL = [
    # Q1: tag OK, mais un seul âge détecté (< 23 seulement)
    ["[tag] classe ii d'angle", "[age] moins de 23 ans",
     "Résidu: 'patients en ages de plus de 20 ans et de'"],
    # Q2: tag OK, âge absent
    ["[tag] classe i d'angle",
     "Résidu: 'patient compris entre 13 et 16 avec une uniquement'"],
    # Q3: tag OK, âge absent
    ["[tag] classe iii d'angle",
     "Résidu: 'patients avec un age superieur a 17 ans et inferieur a 23 ans de'"],
    # Q4: tout OK (enfant + tags fonctionnent)
    ["[tag] classe ii d'angle", "[tag] rétrognathie mandibulaire", "[age] enfant",
     "Résidu: 'patient avec et ou'"],
    # Q5: tag OK, âge absent
    ["[tag] métal",
     "Résidu: 'patient compris entre 20 e 24 ans'"],
]

# detia AVANT : bugs de fusion BETWEEN + classe → squelettique
DETIA = [
    # Q1: "classe 2" → squelettique (bug), âges fusionnés en BETWEEN (bug)
    ["[tag] classe ii squelettique", "[age] Âge BETWEEN(20, 23)",
     "Résidu: 'patients en ages'", "3700ms"],
    # Q2: "entre 13 et 16" → BETWEEN OK (formulation explicite), classe OK
    ["[age] Âge BETWEEN(13, 16)", "[tag] classe i d'angle [uniquement]",
     "Résidu: 'patient avec une de'", "2800ms"],
    # Q3: "classe 3" → squelettique (bug), âges fusionnés en BETWEEN (bug)
    ["[tag] classe iii squelettique", "[age] Âge BETWEEN(17, 23)",
     "Résidu: 'patients avec un et de'", "3500ms"],
    # Q4: rétrognathie OK, mais "classe 2" → squelettique (bug)
    ["[age] Âge < 12", "[tag] rétrognathie mandibulaire", "[tag] classe ii squelettique",
     "Résidu: 'patient avec et ou'", "3400ms"],
    # Q5: "compris entre 20 e 24" → IA gère la faute → BETWEEN OK
    ["[tag] métal", "[age] Âge BETWEEN(20, 24)",
     "Résidu: 'patient'", "2700ms"],
]

# detiabrut AVANT : comme detia mais sans références (pire)
DETIABRUT = [
    # Q1: pas de mapping canonique → "Classe 2" brut, BETWEEN
    ["[tag] Classe 2", "[age] Âge BETWEEN(20, 23)",
     "Résidu: 'patients en ages et'", "2600ms"],
    # Q2: BETWEEN OK, classe brut
    ["[age] Âge BETWEEN(13, 16)", "[tag] Classe D'Angle De 1 [Uniquement]",
     "Résidu: 'patient avec une d angle'", "2400ms"],
    # Q3: "Classe 3" brut, BETWEEN
    ["[tag] Classe 3", "[age] Âge BETWEEN(17, 23)",
     "Résidu: 'patients avec un et de'", "3000ms"],
    # Q4: mapping brut
    ["[tag] Rétrognathie", "[tag] Classe 2", "[age] Âge None None",
     "Résidu: 'patient avec et ou'", "2800ms"],
    # Q5: IA brut gère la faute
    ["[age] Âge BETWEEN(20, 24)",
     "Résidu: 'patient metal compris'", "2000ms"],
]

# trouve AVANT : utilise detall buggé → résultats incorrects
TROUVE = [
    # Q1: manque >20 → trop de patients (classe ii d'angle + <23 seulement)
    ["[tag] classe ii d'angle", "[age] moins de 23 ans",
     "→ 47 patient(s) en 15ms"],
    # Q2: pas d'âge → tous les classe i d'angle
    ["[tag] classe i d'angle",
     "→ 312 patient(s) en 8ms"],
    # Q3: pas d'âge → tous les classe iii d'angle
    ["[tag] classe iii d'angle",
     "→ 312 patient(s) en 5ms"],
    # Q4: tout OK sur le standard
    ["[tag] classe ii d'angle", "[tag] rétrognathie mandibulaire", "[age] enfant",
     "→ 0 patient(s) en 3ms"],
    # Q5: pas d'âge → tous les métal
    ["[tag] métal",
     "→ 89 patient(s) en 4ms"],
]

# trouveid : inchangé (pas d'id dans les questions)
TROUVEID = [
    ["Résidu: 'patients en classe 2 ages de plus de 20 ans et de moins de 23 ans'"],
    ["Résidu: 'patient compris entre 13 et 16 avec une classe d angle de 1 uniquement'"],
    ["Résidu: 'patients avec un age superieur a 17 ans et inferieur a 23 ans de classe 3'"],
    ["Résidu: 'patient enfant avec retrognathie et ou classe 2'"],
    ["Résidu: 'patient metal compris entre 20 e 24 ans'"],
]

# search : ❌ CASSÉ (parser batch rigide + langue "zu")
# Pas de données, module en échec

# ─── Résumé des modules ───

MODULES_RESULTATS = [
    {'module': 'detage',    'succes': True,  'nb': 5, 'duree': 78,    'data': DETAGE},
    {'module': 'detall',    'succes': True,  'nb': 5, 'duree': 115,   'data': DETALL},
    {'module': 'detangles', 'succes': True,  'nb': 5, 'duree': 80,    'data': DETANGLES},
    {'module': 'detcount',  'succes': True,  'nb': 5, 'duree': 62,    'data': DETCOUNT},
    {'module': 'detia',     'succes': True,  'nb': 5, 'duree': 16200, 'data': DETIA},
    {'module': 'detiabrut', 'succes': True,  'nb': 5, 'duree': 13100, 'data': DETIABRUT},
    {'module': 'detid',     'succes': True,  'nb': 5, 'duree': 85,    'data': DETID},
    {'module': 'detmeme',   'succes': True,  'nb': 5, 'duree': 74,    'data': DETMEME},
    {'module': 'dettags',   'succes': True,  'nb': 5, 'duree': 108,   'data': DETTAGS},
    {'module': 'search',    'succes': False, 'nb': 0, 'duree': 1200,  'data': None,
     'erreur': "Parser batch cassé : ordre d'arguments imposé + fichier CSV non trouvé dans tests/ + langue détectée 'zu'"},
    {'module': 'trouve',    'succes': True,  'nb': 5, 'duree': 128,   'data': TROUVE},
    {'module': 'trouveid',  'succes': True,  'nb': 5, 'duree': 90,    'data': TROUVEID},
]

# ─── Évaluations des commentaires (tous KO car bugs présents) ───

EVALUATIONS = [
    {
        'question': QUESTIONS[0]['question'],
        'commentaire': QUESTIONS[0]['commentaire'],
        'statut': 'KO',
        'explication': "Le critère 'plus de 20 ans' n'est pas détecté. Seul 'moins de 23 ans' apparaît. Le bug de tracking en caractères empêche la double détection.",
        'action': None,
    },
    {
        'question': QUESTIONS[1]['question'],
        'commentaire': QUESTIONS[1]['commentaire'],
        'statut': 'KO',
        'explication': "Aucun critère d'âge détecté. La formulation 'compris entre 13 et 16' est absente des patterns ages.csv.",
        'action': None,
    },
    {
        'question': QUESTIONS[2]['question'],
        'commentaire': QUESTIONS[2]['commentaire'],
        'statut': 'KO',
        'explication': "Aucun critère d'âge détecté. 'supérieur à' et 'inférieur à' sont absents des patterns ages.csv. 312 patients retournés au lieu de 12.",
        'action': None,
    },
    {
        'question': QUESTIONS[3]['question'],
        'commentaire': QUESTIONS[3]['commentaire'],
        'statut': 'ACTION',
        'explication': "Le standard (dettags) détecte rétrognathie, mais le module IA (detia) la mappe vers 'classe ii squelettique'. Le résultat dépend du chemin emprunté par search.py.",
        'action': "Vérifier via le web si search.py utilise le chemin standard ou le fallback IA pour cette question.",
    },
    {
        'question': QUESTIONS[4]['question'],
        'commentaire': QUESTIONS[4]['commentaire'],
        'statut': 'KO',
        'explication': "La faute de frappe 'e' au lieu de 'et' empêche la détection BETWEEN par detage. Le critère d'âge est totalement absent.",
        'action': None,
    },
]

# Descriptions des modules
DESCRIPTIONS_MODULES = {
    'detage':    "Détecte l'âge et le sexe des patients recherchés",
    'detall':    "Pipeline complet : détecte tout (tags, âge, angles, meme, id...)",
    'detangles': "Détecte les angles céphalométriques (ANB, SNA, SNB)",
    'detcount':  "Détecte si on veut compter (COUNT) ou lister (LIST)",
    'detid':     "Détecte les identifiants patients (id 10122)",
    'detmeme':   "Détecte les recherches par similarité (même portrait que...)",
    'dettags':   "Détecte les pathologies et leurs adjectifs",
    'detia':     "Détection par intelligence artificielle avec référentiels complets",
    'detiabrut': "Détection par intelligence artificielle sans aide (mode brut)",
    'trouve':    "Recherche standard : detall → jsonsql → lancesql",
    'trouveid':  "Recherche avec résolution des similarités et identifiants",
    'search':    "Pipeline search standard (avec fallback IA automatique)",
}

CATEGORIES = {
    'standard': {
        'titre': 'Détection standard (sans IA)',
        'emoji': '🔍',
        'description': "Ces modules analysent la question mot par mot avec des règles précises.",
        'modules': ['detcount', 'detid', 'detage', 'detangles', 'dettags', 'detmeme', 'detall'],
    },
    'ia': {
        'titre': 'Détection par Intelligence Artificielle',
        'emoji': '🤖',
        'description': "Ces modules envoient la question à une IA qui comprend le langage naturel.",
        'modules': ['detia', 'detiabrut'],
    },
    'recherche': {
        'titre': 'Recherche dans la base de patients',
        'emoji': '🏥',
        'description': "Ces modules cherchent vraiment les patients dans la base de données.",
        'modules': ['trouve', 'trouveid', 'search'],
    },
}


# ═══════════════════════════════════════════════════════════════════════════
# GÉNÉRATION DOCX
# ═══════════════════════════════════════════════════════════════════════════

def _installer_docx():
    """Installe python-docx si nécessaire."""
    try:
        from docx import Document
        return True
    except ImportError:
        print("  Installation de python-docx...")
        proc = subprocess.run(
            [sys.executable, '-m', 'pip', 'install', 'python-docx', '--break-system-packages', '-q'],
            capture_output=True
        )
        return proc.returncode == 0


def _colorer_heading(heading, color):
    for run in heading.runs:
        run.font.color.rgb = color


def _cellule_entete(cell, texte):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texte)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): '1A477A',
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def _cellule_donnee(cell, texte, bold=False):
    from docx.shared import Pt
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(texte)
    run.font.size = Pt(9)
    run.bold = bold


def _cellule_donnee_coloree(cell, texte):
    from docx.shared import Pt, RGBColor
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(texte)
    run.font.size = Pt(8)

    if texte.startswith('[tag]'):
        run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    elif texte.startswith('[age]') or texte.startswith('[sexe]'):
        run.font.color.rgb = RGBColor(0x00, 0x64, 0x00)
    elif texte.startswith('[angle]'):
        run.font.color.rgb = RGBColor(0x80, 0x00, 0x80)
    elif texte.startswith('[id]'):
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x8B)
    elif texte.startswith('[meme]'):
        run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
    elif texte.startswith('[adj]'):
        run.font.color.rgb = RGBColor(0x69, 0x69, 0x69)
    elif texte.startswith('→') or texte.startswith('Résidu'):
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.italic = True
    elif '✗' in texte or 'erreur' in texte.lower():
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    elif texte.startswith('Routage:'):
        run.font.color.rgb = RGBColor(0x00, 0x64, 0x00)
        run.bold = True


def generer_rapport(verbose=False, debug=False):
    """Génère quentin_init.docx — Rapport AVANT corrections."""

    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # ─── Style par défaut ───
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ═══════════════════════════════════════════════════════════════════
    # PAGE DE TITRE
    # ═══════════════════════════════════════════════════════════════════

    for _ in range(4):
        doc.add_paragraph()

    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titre.add_run("📋 Rapport de Tests — AVANT corrections")
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.bold = True

    sous_titre = doc.add_paragraph()
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sous_titre.add_run("KitView Search — Reconstitution état initial")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    infos = doc.add_paragraph()
    infos.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = infos.add_run("Fichier de test : quentin.csv")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    infos2 = doc.add_paragraph()
    infos2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = infos2.add_run("Base de données : base1964.db")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    infos3 = doc.add_paragraph()
    infos3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = infos3.add_run(f"{len(QUESTIONS)} question(s) testée(s) — {len(MODULES_RESULTATS)} module(s)")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Résumé évaluation
    nb_ok = sum(1 for e in EVALUATIONS if e['statut'] == 'OK')
    nb_ko = sum(1 for e in EVALUATIONS if e['statut'] == 'KO')
    nb_action = sum(1 for e in EVALUATIONS if e['statut'] == 'ACTION')

    infos_eval = doc.add_paragraph()
    infos_eval.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = infos_eval.add_run(f"Évaluation : {nb_ok} OK · {nb_ko} KO · {nb_action} ACTION · 0 PASS")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.bold = True

    # Mention reconstitution
    p_reconst = doc.add_paragraph()
    p_reconst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_reconst.add_run("⚠ Ce rapport est une reconstitution des résultats AVANT corrections")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
    run.italic = True

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("20/02/2026 — Date des tests originaux de Quentin")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # LÉGENDE DES COULEURS
    # ═══════════════════════════════════════════════════════════════════

    h = doc.add_heading("🎨  Légende des couleurs", level=1)
    _colorer_heading(h, RGBColor(0x1A, 0x47, 0x7A))

    p = doc.add_paragraph()
    run = p.add_run("Les résultats dans les tableaux sont colorés selon le type de critère détecté :")
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    legende = [
        ("[tag]",          "Pathologie / tag médical",          RGBColor(0x8B, 0x00, 0x00)),
        ("[age] / [sexe]", "Critère d'âge ou de sexe",          RGBColor(0x00, 0x64, 0x00)),
        ("[angle]",        "Angle céphalométrique (ANB, SNA…)", RGBColor(0x80, 0x00, 0x80)),
        ("[id]",           "Identifiant patient",               RGBColor(0x00, 0x00, 0x8B)),
        ("[meme]",         "Similarité (même X que…)",          RGBColor(0xFF, 0x8C, 0x00)),
        ("[adj]",          "Adjectif qualificatif",             RGBColor(0x69, 0x69, 0x69)),
        ("→ / Résidu",     "Résultat final ou résidu",          RGBColor(0x55, 0x55, 0x55)),
        ("✗ / erreur",     "Erreur ou échec",                   RGBColor(0xCC, 0x00, 0x00)),
        ("✅ OK",          "Commentaire résolu par la détection",    RGBColor(0x00, 0x80, 0x00)),
        ("❌ KO",          "Problème persistant malgré la détection", RGBColor(0xCC, 0x00, 0x00)),
        ("🔧 ACTION",      "Vérification manuelle requise",          RGBColor(0xFF, 0x8C, 0x00)),
    ]

    table_leg = doc.add_table(rows=1, cols=3)
    table_leg.style = 'Table Grid'
    _cellule_entete(table_leg.rows[0].cells[0], "Préfixe")
    _cellule_entete(table_leg.rows[0].cells[1], "Signification")
    _cellule_entete(table_leg.rows[0].cells[2], "Exemple")

    for prefixe, signification, couleur in legende:
        row = table_leg.add_row()
        cell_p = row.cells[0]
        cell_p.text = ''
        p_run = cell_p.paragraphs[0].add_run(prefixe)
        p_run.font.size = Pt(9)
        p_run.font.color.rgb = couleur
        p_run.bold = True
        _cellule_donnee(row.cells[1], signification)
        _cellule_donnee(row.cells[2], "")

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # SOMMAIRE
    # ═══════════════════════════════════════════════════════════════════

    h = doc.add_heading("Sommaire", level=1)
    _colorer_heading(h, RGBColor(0x1A, 0x47, 0x7A))

    modules_par_nom = {r['module']: r for r in MODULES_RESULTATS}

    for cat_key, cat in CATEGORIES.items():
        modules_cat = [m for m in cat['modules'] if m in modules_par_nom]
        if not modules_cat:
            continue
        p = doc.add_paragraph()
        run = p.add_run(f"{cat['emoji']}  {cat['titre']}")
        run.bold = True
        run.font.size = Pt(12)
        for m in modules_cat:
            r = modules_par_nom[m]
            p2 = doc.add_paragraph(style='List Bullet')
            status = "✓" if r['succes'] else "✗"
            p2.add_run(f"{status}  {m}.py — {DESCRIPTIONS_MODULES.get(m, '')}")

    p = doc.add_paragraph()
    run = p.add_run("📝  Évaluation des commentaires")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # QUESTIONS TESTÉES
    # ═══════════════════════════════════════════════════════════════════

    h = doc.add_heading("Questions testées", level=1)
    _colorer_heading(h, RGBColor(0x1A, 0x47, 0x7A))

    p = doc.add_paragraph()
    p.add_run("Le fichier ").font.size = Pt(11)
    run = p.add_run("quentin.csv")
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(f" contient {len(QUESTIONS)} question(s) :").font.size = Pt(11)

    for i, q in enumerate(QUESTIONS, 1):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"Q{i}: ")
        run.bold = True
        p.add_run(q['question'])
        if q['commentaire']:
            run_c = p.add_run(f"\n    💬 {q['commentaire']}")
            run_c.italic = True
            run_c.font.size = Pt(9)
            run_c.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # ÉVALUATION DES COMMENTAIRES
    # ═══════════════════════════════════════════════════════════════════

    h = doc.add_heading("📝  Évaluation des commentaires", level=1)
    _colorer_heading(h, RGBColor(0x1A, 0x47, 0x7A))

    p = doc.add_paragraph()
    run = p.add_run(
        "Chaque commentaire du testeur a été évalué contre les résultats du module detall "
        "(pipeline complet de détection). "
        "État AVANT corrections — les bugs sont encore présents."
    )
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    table_eval = doc.add_table(rows=1, cols=4)
    table_eval.style = 'Table Grid'
    _cellule_entete(table_eval.rows[0].cells[0], "Question")
    _cellule_entete(table_eval.rows[0].cells[1], "Commentaire")
    _cellule_entete(table_eval.rows[0].cells[2], "Statut")
    _cellule_entete(table_eval.rows[0].cells[3], "Explication / Action")

    for ev in EVALUATIONS:
        row_ev = table_eval.add_row()

        q_text = ev['question']
        if len(q_text) > 45:
            q_text = q_text[:42] + "..."
        _cellule_donnee(row_ev.cells[0], q_text, bold=True)

        comm = ev['commentaire']
        if len(comm) > 60:
            comm = comm[:57] + "..."
        _cellule_donnee(row_ev.cells[1], comm)

        statut = ev['statut']
        cell_statut = row_ev.cells[2]
        cell_statut.text = ''
        p_s = cell_statut.paragraphs[0]
        emoji = {'OK': '✅', 'KO': '❌', 'ACTION': '🔧'}.get(statut, '❓')
        run_s = p_s.add_run(f"{emoji} {statut}")
        run_s.bold = True
        run_s.font.size = Pt(9)
        couleurs_statut = {
            'OK': RGBColor(0x00, 0x80, 0x00),
            'KO': RGBColor(0xCC, 0x00, 0x00),
            'ACTION': RGBColor(0xFF, 0x8C, 0x00),
        }
        run_s.font.color.rgb = couleurs_statut.get(statut, RGBColor(0x00, 0x00, 0x00))

        cell_expl = row_ev.cells[3]
        cell_expl.text = ''
        p_e = cell_expl.paragraphs[0]
        run_e = p_e.add_run(ev.get('explication', ''))
        run_e.font.size = Pt(8)
        if ev.get('action'):
            run_a = p_e.add_run(f"\n📋 {ev['action']}")
            run_a.font.size = Pt(8)
            run_a.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
            run_a.bold = True

    for row in table_eval.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.size is None:
                        run.font.size = Pt(8)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(f"Bilan : {nb_ok} OK · {nb_ko} KO · {nb_action} ACTION · 0 PASS")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # RÉSULTATS PAR CATÉGORIE
    # ═══════════════════════════════════════════════════════════════════

    for cat_key, cat in CATEGORIES.items():
        modules_cat = [m for m in cat['modules'] if m in modules_par_nom]
        if not modules_cat:
            continue

        h = doc.add_heading(f"{cat['emoji']}  {cat['titre']}", level=1)
        _colorer_heading(h, RGBColor(0x1A, 0x47, 0x7A))

        p = doc.add_paragraph()
        run = p.add_run(cat['description'])
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        for module in modules_cat:
            r = modules_par_nom[module]

            h2 = doc.add_heading(f"{module}.py", level=2)
            _colorer_heading(h2, RGBColor(0x2E, 0x75, 0xB6))

            desc = DESCRIPTIONS_MODULES.get(module, '')
            if desc:
                p = doc.add_paragraph()
                run = p.add_run(desc)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            if r['succes']:
                p = doc.add_paragraph()
                run = p.add_run(f"✅ Succès — {r['nb']} question(s) en {r['duree']}ms")
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                run.bold = True
            else:
                p = doc.add_paragraph()
                erreur = r.get('erreur', 'Erreur inconnue')
                run = p.add_run(f"❌ Échec — {erreur}")
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                run.bold = True
                continue  # Pas de tableau pour un module en échec

            data = r.get('data')
            if data:
                max_l = max(len(lignes) for lignes in data)
                nb_cols = 1 + max_l
                table = doc.add_table(rows=1, cols=nb_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.LEFT

                hdr = table.rows[0]
                _cellule_entete(hdr.cells[0], "Question")
                for ci in range(max_l):
                    _cellule_entete(hdr.cells[1 + ci], f"L{ci + 1}")

                for qi, lignes in enumerate(data):
                    row_t = table.add_row()
                    q_text = QUESTIONS[qi]['question']
                    if len(q_text) > 50:
                        q_text = q_text[:47] + "..."
                    _cellule_donnee(row_t.cells[0], q_text, bold=True)

                    for ci in range(max_l):
                        texte = lignes[ci] if ci < len(lignes) else ''
                        _cellule_donnee_coloree(row_t.cells[1 + ci], texte)

                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)

                doc.add_paragraph()

        doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # SYNTHÈSE GLOBALE
    # ═══════════════════════════════════════════════════════════════════

    h = doc.add_heading("📊  Synthèse globale", level=1)
    _colorer_heading(h, RGBColor(0x1A, 0x47, 0x7A))

    nb_ok_mod = sum(1 for r in MODULES_RESULTATS if r['succes'])
    nb_ko_mod = sum(1 for r in MODULES_RESULTATS if not r['succes'])
    duree_totale = sum(r['duree'] for r in MODULES_RESULTATS)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    _cellule_entete(table.rows[0].cells[0], "Module")
    _cellule_entete(table.rows[0].cells[1], "Statut")
    _cellule_entete(table.rows[0].cells[2], "Questions")
    _cellule_entete(table.rows[0].cells[3], "Durée")

    for r in MODULES_RESULTATS:
        row_s = table.add_row()
        _cellule_donnee(row_s.cells[0], r['module'], bold=True)
        if r['succes']:
            _cellule_donnee(row_s.cells[1], "✅ OK")
        else:
            _cellule_donnee(row_s.cells[1], "❌ KO")
        _cellule_donnee(row_s.cells[2], str(r['nb']))
        _cellule_donnee(row_s.cells[3], f"{r['duree']}ms")

    row_total = table.add_row()
    _cellule_donnee(row_total.cells[0], "TOTAL", bold=True)
    _cellule_donnee(row_total.cells[1], f"{nb_ok_mod} ✅  {nb_ko_mod} ❌", bold=True)
    _cellule_donnee(row_total.cells[2], str(len(QUESTIONS)), bold=True)
    _cellule_donnee(row_total.cells[3], f"{duree_totale}ms", bold=True)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # Conclusion modules
    p = doc.add_paragraph()
    run = p.add_run(f"⚠ {nb_ko_mod} module(s) en échec sur {len(MODULES_RESULTATS)}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xCC, 0x80, 0x00)
    run.bold = True

    # Conclusion évaluation
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"❌ {nb_ko} problème(s) persistant(s), {nb_action} vérification(s) manuelle(s)")
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.size = Pt(12)
    run.bold = True

    # ─── Sauvegarde ───
    script_dir = Path(__file__).parent
    rapport_path = script_dir / "quentin_init.docx"
    doc.save(str(rapport_path))

    return rapport_path


# ═══════════════════════════════════════════════════════════════════════════
# AIDE CLI
# ═══════════════════════════════════════════════════════════════════════════

def afficher_aide():
    """Affiche l'aide."""
    print("Usage:")
    print(f"  python {__pgm__} all          Génère quentin_init.docx")
    print(f"  python {__pgm__} all -v       Mode verbose")
    print(f"  python {__pgm__} all -d       Mode debug")
    print()
    print("Options:")
    print("  -v   Mode verbose")
    print("  -d   Mode debug")
    print()
    print("Ce script génère un rapport DOCX reconstituant les résultats")
    print("AVANT les corrections de bugs identifiés par Quentin.")
    print()
    print("Bugs reconstitués :")
    print("  - detage.py : tracking en caractères au lieu de mots")
    print("  - ages.csv  : patterns manquants")
    print("  - detia.py  : fusion BETWEEN + classe → squelettique")
    print("  - search.py : parser batch cassé + langue 'zu'")
    print()


# ═══════════════════════════════════════════════════════════════════════════
# POINT D'ENTRÉE
# ═══════════════════════════════════════════════════════════════════════════

def main():
    """Point d'entrée CLI."""
    print(f"╔════════════════════════════════════════════════════════════════")
    print(f"║ {__pgm__} V{__version__} - {__date__}")
    print(f"║ Génération du rapport quentin_init.docx (AVANT corrections)")
    print(f"╚════════════════════════════════════════════════════════════════")
    print()

    if len(sys.argv) < 2:
        afficher_aide()
        return 1

    verbose = '-v' in sys.argv or '--verbose' in sys.argv
    debug = '-d' in sys.argv or '--debug' in sys.argv

    commande = None
    for arg in sys.argv[1:]:
        if not arg.startswith('-'):
            commande = arg
            break

    if commande != 'all':
        afficher_aide()
        return 1

    # Installer python-docx si nécessaire
    if not _installer_docx():
        print("ERREUR : Impossible d'installer python-docx")
        return 1

    print("Génération du rapport...")
    print()

    rapport_path = generer_rapport(verbose=verbose, debug=debug)

    print()
    print(f"✓ Rapport généré : {os.path.abspath(rapport_path)}")
    print()

    # Résumé
    nb_ko = sum(1 for e in EVALUATIONS if e['statut'] == 'KO')
    nb_action = sum(1 for e in EVALUATIONS if e['statut'] == 'ACTION')
    nb_echec = sum(1 for r in MODULES_RESULTATS if not r['succes'])
    print(f"  {len(MODULES_RESULTATS)} modules testés, {nb_echec} en échec")
    print(f"  {len(EVALUATIONS)} commentaires évalués : {nb_ko} KO, {nb_action} ACTION")
    print()

    return 0


if __name__ == '__main__':
    sys.exit(main())
